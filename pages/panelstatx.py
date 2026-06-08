"""
apps/panelstatx.py
══════════════════════════════════════════════════════════════════════
Bayantx360 Suite — PanelStatX
Panel Data Econometrics Engine
══════════════════════════════════════════════════════════════════════

Changes from standalone version:
  • Auth + credit engine → shared/auth.py  (removed local duplicates)
  • CSS → shared/theme.py                  (removed local CSS block)
  • Single secret: BAYANTX_SHEET_ID        (was SHEET_ID)
  • Trial gate: is_free_trial flag          (standardised — no fake credits)
  • Back-to-suite navigation in sidebar

All statistical logic (OLS, FE, RE, FD, diagnostics, DOCX) is 100%
preserved from the original app__28_.py.
"""

import streamlit as st
import pandas as pd
import numpy as np
import plotly.express as px
import plotly.graph_objects as go
import requests
import io
import warnings
import sys, os
warnings.filterwarnings("ignore")

# ── Path resolution ────────────────────────────────────────────────────────────

sys.path.insert(0, os.path.join(os.path.dirname(os.path.abspath(__file__)), ".."))
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from shared.auth import (
    init_session_state,
    refresh_credits,
    handle_credit_deduction,
    render_credit_hud,
    can_use_premium,
    is_trial,
    sign_out,
)
from shared.theme import apply_suite_css, apply_theme, PLOTLY_THEME, render_locked_banner

# ── Page config ────────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="PanelStatX · Bayantx360 Suite",
    page_icon="📐",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ── Init ───────────────────────────────────────────────────────────────────────
init_session_state()
apply_suite_css()

# ── Auth guard: redirect to suite home if not authenticated ───────────────────
if not st.session_state.get("access_granted"):
    st.switch_page(st.session_state["_home_page"])
    st.stop()

# Refresh live credit balance on every page load (paid users only)
refresh_credits()

# ── Per-app session keys ───────────────────────────────────────────────────────
for key, default in [
    ("df", None), ("results", None), ("ai_explanation", ""),
    ("model_type", "Fixed Effects (Two-Way)"),
]:
    if key not in st.session_state:
        st.session_state[key] = default


# ═══════════════════════════════════════════════════════════════════════════════
# STATISTICAL FUNCTIONS  (100% preserved from original)
# ═══════════════════════════════════════════════════════════════════════════════

def generate_demo_panel():
    np.random.seed(42)
    n_entities, n_periods = 30, 10
    entities = [f"Entity_{i:02d}" for i in range(1, n_entities + 1)]
    years = list(range(2014, 2014 + n_periods))
    rows = []
    for e in entities:
        fe = np.random.randn()
        for y in years:
            te = 0.05 * (y - 2014)
            x1 = np.random.randn() + fe * 0.3
            x2 = np.random.uniform(0, 10)
            x3 = np.random.choice([0, 1], p=[0.6, 0.4])
            y_val = 2 + 0.8 * x1 - 0.4 * x2 + 1.2 * x3 + fe + te + np.random.randn() * 0.5
            rows.append({"entity": e, "year": y, "y": round(y_val, 4),
                         "x1": round(x1, 4), "x2": round(x2, 4), "x3": int(x3)})
    return pd.DataFrame(rows)


def run_ols(df, y_col, x_cols):
    from numpy.linalg import lstsq
    from scipy import stats as sc_stats
    X = np.column_stack([np.ones(len(df))] + [df[c].values for c in x_cols])
    y = df[y_col].values
    coeffs, _, _, _ = lstsq(X, y, rcond=None)
    y_hat = X @ coeffs
    resid = y - y_hat
    n, k = X.shape
    k_vars = k - 1
    dof = n - k
    s2 = np.sum(resid**2) / max(dof, 1)
    cov = s2 * np.linalg.inv(X.T @ X)
    se = np.sqrt(np.diag(cov))
    t_stats = coeffs / se
    p_vals = 2 * sc_stats.t.sf(np.abs(t_stats), df=dof)
    ss_tot = np.sum((y - y.mean())**2)
    ss_res = np.sum(resid**2)
    ss_reg = ss_tot - ss_res
    r2 = 1 - ss_res / max(ss_tot, 1e-12)
    r2_adj = 1 - (1 - r2) * (n - 1) / max(dof, 1)
    f_stat = (ss_reg / max(k_vars, 1)) / (ss_res / max(dof, 1))
    f_p = 1 - sc_stats.f.cdf(f_stat, dfn=k_vars, dfd=dof)
    names = ["const"] + list(x_cols)
    result_df = pd.DataFrame({"Variable": names, "Coeff": coeffs, "Std_Err": se,
                               "t_stat": t_stats, "p_value": p_vals})
    stats = {"R2": r2, "R2_adj": r2_adj, "N": n, "k": k_vars,
             "AIC": n * np.log(ss_res / max(n, 1)) + 2 * k,
             "BIC": n * np.log(ss_res / max(n, 1)) + k * np.log(n),
             "F_stat": f_stat, "F_p": f_p}
    return result_df, resid, y_hat, stats, cov


def run_within(df, y_col, x_cols, entity_col, time_col):
    from scipy import stats as sc_stats
    panel = df.copy()
    for col in [y_col] + list(x_cols):
        entity_means = panel.groupby(entity_col)[col].transform("mean")
        time_means = panel.groupby(time_col)[col].transform("mean")
        grand_mean = panel[col].mean()
        panel[col + "_dm"] = panel[col] - entity_means - time_means + grand_mean
    y_dm = panel[y_col + "_dm"].values
    X_dm = np.column_stack([panel[c + "_dm"].values for c in x_cols])
    from numpy.linalg import lstsq
    coeffs, _, _, _ = lstsq(X_dm, y_dm, rcond=None)
    y_hat_dm = X_dm @ coeffs
    resid = y_dm - y_hat_dm
    n, k = X_dm.shape
    dof = n - k - df[entity_col].nunique() - df[time_col].nunique() + 1
    if dof <= 0:
        dof = max(1, n - k)
    s2 = np.sum(resid**2) / dof
    cov = s2 * np.linalg.inv(X_dm.T @ X_dm)
    se = np.sqrt(np.diag(cov))
    t_stats = coeffs / se
    p_vals = 2 * sc_stats.t.sf(np.abs(t_stats), df=dof)
    ss_tot = np.sum((y_dm - y_dm.mean())**2)
    ss_res = np.sum(resid**2)
    ss_reg = ss_tot - ss_res
    r2 = max(0, 1 - ss_res / max(ss_tot, 1e-12))
    r2_adj = max(0, 1 - (1 - r2) * (n - 1) / max(dof, 1))
    f_stat = (ss_reg / max(k, 1)) / (ss_res / max(dof, 1))
    f_p = 1 - sc_stats.f.cdf(f_stat, dfn=k, dfd=dof)
    result_df = pd.DataFrame({"Variable": list(x_cols), "Coeff": coeffs,
                               "Std_Err": se, "t_stat": t_stats, "p_value": p_vals})
    stats = {"R2": r2, "R2_adj": r2_adj, "N": n, "k": k,
             "AIC": n * np.log(max(ss_res, 1e-10) / n) + 2 * k,
             "BIC": n * np.log(max(ss_res, 1e-10) / n) + k * np.log(n),
             "F_stat": f_stat, "F_p": f_p}
    return result_df, resid, y_hat_dm, stats, cov


def run_re(df, y_col, x_cols, entity_col, time_col):
    from numpy.linalg import lstsq
    from scipy import stats as sc_stats
    panel = df.copy().sort_values([entity_col, time_col])
    n_entities = panel[entity_col].nunique()
    T = panel[time_col].nunique()
    N = len(panel)
    k = len(x_cols)
    result_fe, resid_fe, _, stats_fe, _ = run_within(panel, y_col, x_cols, entity_col, time_col)
    dof_fe = max(N - k - n_entities, 1)
    sigma_e2 = np.sum(resid_fe ** 2) / dof_fe
    grp = panel.groupby(entity_col)[[y_col] + list(x_cols)].mean().reset_index()
    y_b = grp[y_col].values
    X_b = np.column_stack([np.ones(n_entities)] + [grp[c].values for c in x_cols])
    b_coeffs, _, _, _ = lstsq(X_b, y_b, rcond=None)
    resid_b = y_b - X_b @ b_coeffs
    sigma_b2 = max(0.0, np.sum(resid_b ** 2) / max(n_entities - k - 1, 1) - sigma_e2 / T)
    sigma_u2 = sigma_b2
    theta = 1.0 - np.sqrt(sigma_e2 / max(T * sigma_u2 + sigma_e2, 1e-12))
    panel2 = panel.copy()
    for col in [y_col] + list(x_cols):
        entity_mean = panel2.groupby(entity_col)[col].transform("mean")
        panel2[col + "_qd"] = panel2[col] - theta * entity_mean
    y_qd = panel2[y_col + "_qd"].values
    X_qd = np.column_stack([np.ones(N)] + [panel2[c + "_qd"].values for c in x_cols])
    coeffs, _, _, _ = lstsq(X_qd, y_qd, rcond=None)
    y_hat = X_qd @ coeffs
    resid = y_qd - y_hat
    dof = max(N - k - 1, 1)
    s2 = np.sum(resid ** 2) / dof
    cov = s2 * np.linalg.inv(X_qd.T @ X_qd)
    se = np.sqrt(np.diag(cov))
    t_stats = coeffs / se
    p_vals = 2 * sc_stats.t.sf(np.abs(t_stats), df=dof)
    ss_tot = np.sum((y_qd - y_qd.mean()) ** 2)
    ss_res = np.sum(resid ** 2)
    ss_reg = ss_tot - ss_res
    r2 = max(0.0, 1 - ss_res / max(ss_tot, 1e-12))
    r2_adj = max(0.0, 1 - (1 - r2) * (N - 1) / dof)
    f_stat = (ss_reg / max(k, 1)) / (ss_res / dof)
    f_p = 1 - sc_stats.f.cdf(f_stat, dfn=k, dfd=dof)
    names = ["const"] + list(x_cols)
    result_df = pd.DataFrame({"Variable": names, "Coeff": coeffs, "Std_Err": se,
                               "t_stat": t_stats, "p_value": p_vals})
    stats = {
        "R2": r2, "R2_adj": r2_adj, "N": N, "k": k,
        "AIC": N * np.log(max(ss_res, 1e-10) / N) + 2 * (k + 1),
        "BIC": N * np.log(max(ss_res, 1e-10) / N) + (k + 1) * np.log(N),
        "F_stat": f_stat, "F_p": f_p,
        "sigma_u2": sigma_u2, "sigma_e2": sigma_e2, "theta": theta,
    }
    return result_df, resid, y_hat, stats, cov


def run_fd(df, y_col, x_cols, entity_col, time_col):
    panel = df.sort_values([entity_col, time_col]).copy()
    fd = panel.groupby(entity_col)[[y_col] + list(x_cols)].diff().dropna()
    return run_ols(fd, y_col, x_cols)


def breusch_pagan_test(resid, X):
    from scipy import stats as sc_stats
    from numpy.linalg import lstsq
    resid = np.asarray(resid, dtype=float)
    e2 = resid ** 2
    coeffs_aux, _, _, _ = lstsq(X, e2, rcond=None)
    e2_hat = X @ coeffs_aux
    ss_tot_aux = np.sum((e2 - e2.mean()) ** 2)
    ss_res_aux = np.sum((e2 - e2_hat) ** 2)
    r2_aux = max(0.0, 1 - ss_res_aux / max(ss_tot_aux, 1e-12))
    n = len(resid)
    k = X.shape[1] - 1
    bp_stat = n * r2_aux
    bp_p = 1 - sc_stats.chi2.cdf(bp_stat, df=k)
    return bp_stat, bp_p, k


def hausman_test(fe_coef, re_coef, fe_vcov, re_vcov):
    diff = fe_coef - re_coef
    diff_vcov = fe_vcov - re_vcov
    try:
        stat = float(diff @ np.linalg.inv(diff_vcov) @ diff)
        df_h = len(diff)
        from scipy import stats as sc_stats
        p = 1 - sc_stats.chi2.cdf(stat, df_h)
        return stat, p, df_h
    except Exception:
        return None, None, None


def significance_stars(p):
    if p < 0.001: return "***"
    if p < 0.01:  return "**"
    if p < 0.05:  return "*"
    if p < 0.1:   return "·"
    return ""


def call_openai(system_prompt, user_prompt):
    try:
        api_key = st.secrets["OPENAI_API_KEY"]
    except Exception:
        return "OpenAI API key not configured or exhausted."
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {api_key}",
    }
    payload = {
        "model": "gpt-4o",
        "max_tokens": 1200,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user",   "content": user_prompt},
        ],
    }
    try:
        resp = requests.post(
            "https://api.openai.com/v1/chat/completions",
            headers=headers, json=payload, timeout=30,
        )
        data = resp.json()
        if "choices" in data:
            return data["choices"][0]["message"]["content"]
        return f"API error: {data.get('error', {}).get('message', str(data))}"
    except Exception as e:
        return f"Request failed: {e}"


def build_docx_report(res, model_type, ai_explanation=""):
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import datetime

    doc = Document()
    for section in doc.sections:
        section.page_width  = Inches(8.5)
        section.page_height = Inches(11)
        section.left_margin = section.right_margin = Inches(1.0)
        section.top_margin  = section.bottom_margin = Inches(1.0)

    def shade_cell(cell, hex_color):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        tcPr.append(shd)

    def set_cell_border(cell, **kwargs):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            if edge in kwargs:
                tag = OxmlElement(f'w:{edge}')
                tag.set(qn('w:val'),   kwargs[edge].get('val', 'single'))
                tag.set(qn('w:sz'),    str(kwargs[edge].get('sz', 4)))
                tag.set(qn('w:color'), kwargs[edge].get('color', '000000'))
                tcBorders.append(tag)
        tcPr.append(tcBorders)

    DARK_BG   = "0A0C10"
    TEAL_RGB  = RGBColor(0x00, 0xC8, 0xB0)
    GRAY_RGB  = RGBColor(0x6B, 0x7A, 0x9A)
    TEXT_DARK = RGBColor(0x1A, 0x20, 0x35)

    # Cover
    cover_tbl = doc.add_table(rows=1, cols=1)
    cover_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cover_cell = cover_tbl.rows[0].cells[0]
    shade_cell(cover_cell, DARK_BG)

    cp = cover_cell.add_paragraph()
    cp.paragraph_format.space_before = Pt(18)
    cp.paragraph_format.space_after  = Pt(2)
    r1 = cp.add_run("⬡  PanelStatX · Bayantx360 Suite")
    r1.font.size = Pt(22); r1.font.bold = True
    r1.font.name = "Arial"; r1.font.color.rgb = TEAL_RGB

    cp2 = cover_cell.add_paragraph()
    cp2.paragraph_format.space_after = Pt(4)
    r2 = cp2.add_run("Panel Regression Analysis Report")
    r2.font.size = Pt(14); r2.font.name = "Arial"
    r2.font.color.rgb = RGBColor(0xE2, 0xE8, 0xF4)

    cp3 = cover_cell.add_paragraph()
    cp3.paragraph_format.space_after = Pt(18)
    now_str = datetime.datetime.now().strftime("%d %B %Y, %H:%M")
    r3 = cp3.add_run(f"Generated: {now_str}   ·   Estimator: {model_type}")
    r3.font.size = Pt(9); r3.font.name = "Arial"
    r3.font.color.rgb = GRAY_RGB

    doc.add_paragraph()
    result_df = res["result_df"]
    stats     = res["stats"]
    resid_raw = np.asarray(res["resid"], dtype=float)
    resid_arr = resid_raw[np.isfinite(resid_raw)]

    meta_items = [
        ("Dependent Variable", res["y_col"]),
        ("Independent Variables", ", ".join(res["x_cols"])),
        ("Entity Column", res["entity_col"]),
        ("Time Column", res["time_col"]),
        ("Observations (N)", f"{stats['N']:,}"),
        ("Variables (k)", str(stats["k"])),
    ]
    meta_tbl = doc.add_table(rows=len(meta_items), cols=2)
    meta_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, (label, value) in enumerate(meta_items):
        row = meta_tbl.rows[i]
        lc, vc = row.cells[0], row.cells[1]
        for cell in [lc, vc]:
            for edge in ('top', 'bottom', 'left', 'right'):
                set_cell_border(cell, **{edge: {'val': 'none', 'sz': 0, 'color': 'FFFFFF'}})
        lp = lc.add_paragraph(label)
        lp.runs[0].font.size = Pt(9); lp.runs[0].font.name = "Arial"
        lp.runs[0].bold = True; lp.runs[0].font.color.rgb = GRAY_RGB
        vp = vc.add_paragraph(value)
        vp.runs[0].font.size = Pt(9); vp.runs[0].font.name = "Arial"
        vp.runs[0].font.color.rgb = TEXT_DARK

    doc.add_page_break()

    # Section 1: Fit Statistics
    s1 = doc.add_paragraph()
    s1.paragraph_format.space_before = Pt(0)
    r = s1.add_run("Table 1.  Model Fit Statistics")
    r.bold = True; r.font.size = Pt(12); r.font.name = "Arial"
    r.font.color.rgb = TEXT_DARK

    fit_rows = [
        ("R²",                           f"{stats['R2']:.4f}"),
        ("Adjusted R²",                  f"{stats['R2_adj']:.4f}"),
        ("F-statistic",                  f"{stats.get('F_stat', float('nan')):.4f}"),
        ("Prob (F-statistic)",           f"{stats.get('F_p', float('nan')):.4f}"),
        ("Akaike Info. Criterion (AIC)", f"{stats['AIC']:.2f}"),
        ("Bayesian Info. Criterion (BIC)", f"{stats['BIC']:.2f}"),
        ("No. of Observations",          f"{stats['N']:,}"),
        ("No. of Regressors (k)",        f"{stats['k']}"),
    ]
    if stats.get("sigma_u2") is not None:
        fit_rows += [
            ("Variance: Between-Entity (σ²ᵤ)", f"{stats['sigma_u2']:.6f}"),
            ("Variance: Within-Entity (σ²ₑ)",  f"{stats['sigma_e2']:.6f}"),
            ("GLS Theta (θ)",                   f"{stats['theta']:.4f}"),
        ]

    THICK = {'val': 'single', 'sz': 12, 'color': '1A2035'}
    THIN  = {'val': 'single', 'sz':  4, 'color': '9AA3BE'}
    NO    = None

    def clear_tbl_borders(tbl):
        tbl_el = tbl._tbl
        tblPr = tbl_el.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            tbl_el.insert(0, tblPr)
        for ex in tblPr.findall(qn("w:tblBorders")):
            tblPr.remove(ex)
        tblBorders = OxmlElement("w:tblBorders")
        for edge in ("top","left","bottom","right","insideH","insideV"):
            tag = OxmlElement(f"w:{edge}")
            tag.set(qn("w:val"), "none")
            tag.set(qn("w:sz"), "0")
            tag.set(qn("w:color"), "FFFFFF")
            tblBorders.append(tag)
        tblPr.append(tblBorders)

    def set_academic_border(cell, top=None, bottom=None):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for edge, spec in [('top', top), ('bottom', bottom), ('left', NO), ('right', NO)]:
            tag = OxmlElement(f'w:{edge}')
            if spec:
                tag.set(qn('w:val'),   spec.get('val', 'single'))
                tag.set(qn('w:sz'),    str(spec.get('sz', 4)))
                tag.set(qn('w:color'), spec.get('color', '000000'))
            else:
                tag.set(qn('w:val'), 'none')
                tag.set(qn('w:sz'),  '0')
                tag.set(qn('w:color'), 'FFFFFF')
            tcBorders.append(tag)
        tcPr.append(tcBorders)

    fit_tbl = doc.add_table(rows=len(fit_rows) + 2, cols=2)
    clear_tbl_borders(fit_tbl)
    for j, row_data in enumerate(fit_rows):
        i = j + 1
        row = fit_tbl.rows[i]
        lc, vc = row.cells[0], row.cells[1]
        lc.width = Inches(3.8); vc.width = Inches(2.7)
        top_b = THICK if i == 1 else (THIN if i == 2 else NO)
        bot_b = THICK if i == len(fit_rows) else NO
        set_academic_border(lc, top=top_b, bottom=bot_b)
        set_academic_border(vc, top=top_b, bottom=bot_b)
        lp = lc.paragraphs[0]; vp = vc.paragraphs[0]
        lp.paragraph_format.space_before = lp.paragraph_format.space_after = Pt(3)
        vp.paragraph_format.space_before = vp.paragraph_format.space_after = Pt(3)
        vp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        lr = lp.add_run(row_data[0]); vr = vp.add_run(row_data[1])
        for run in [lr, vr]:
            run.font.size = Pt(9); run.font.name = "Arial"
            run.font.color.rgb = TEXT_DARK

    doc.add_paragraph()

    # Section 2: Coefficient Table
    s2 = doc.add_paragraph()
    r2h = s2.add_run("Table 2.  Coefficient Estimates")
    r2h.bold = True; r2h.font.size = Pt(12); r2h.font.name = "Arial"
    r2h.font.color.rgb = TEXT_DARK

    col_headers = ["Variable", "Coeff.", "Std. Err.", "t-stat", "p-value", "Sig."]
    n_rows = len(result_df) + 2
    coef_tbl = doc.add_table(rows=n_rows, cols=len(col_headers))
    clear_tbl_borders(coef_tbl)

    hdr_row = coef_tbl.rows[1]
    for j, h in enumerate(col_headers):
        c = hdr_row.cells[j]
        set_academic_border(c, top=THICK, bottom=THIN)
        p = c.paragraphs[0]
        p.paragraph_format.space_before = p.paragraph_format.space_after = Pt(3)
        r_h = p.add_run(h)
        r_h.font.size = Pt(9); r_h.bold = True; r_h.font.name = "Arial"
        r_h.font.color.rgb = TEXT_DARK
        if j > 0:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    for i, (_, data_row) in enumerate(result_df.iterrows()):
        row = coef_tbl.rows[i + 2]
        bot_b = THICK if i == len(result_df) - 1 else NO
        values = [
            data_row["Variable"],
            f"{data_row['Coeff']:.4f}",
            f"{data_row['Std_Err']:.4f}",
            f"{data_row['t_stat']:.3f}",
            f"{data_row['p_value']:.4f}",
            significance_stars(data_row["p_value"]),
        ]
        for j, val in enumerate(values):
            c = row.cells[j]
            set_academic_border(c, bottom=bot_b)
            p = c.paragraphs[0]
            p.paragraph_format.space_before = p.paragraph_format.space_after = Pt(3)
            r_v = p.add_run(val)
            r_v.font.size = Pt(9); r_v.font.name = "Arial"
            r_v.font.color.rgb = TEXT_DARK
            if j > 0:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    note_p = doc.add_paragraph()
    note_p.paragraph_format.space_before = Pt(4)
    note_r = note_p.add_run("Note: · p<0.1  * p<0.05  ** p<0.01  *** p<0.001")
    note_r.italic = True; note_r.font.size = Pt(8); note_r.font.name = "Arial"
    note_r.font.color.rgb = GRAY_RGB

    doc.add_paragraph()

    # Section 3: Diagnostics
    from scipy import stats as sc_stats
    jb_stat, jb_p = sc_stats.jarque_bera(resid_arr)
    dw = np.sum(np.diff(resid_arr)**2) / max(np.sum(resid_arr**2), 1e-10)
    skew_v = sc_stats.skew(resid_arr)
    kurt_v = sc_stats.kurtosis(resid_arr)
    bp_stat = stats.get("BP_stat")
    bp_p_v  = stats.get("BP_p")

    s3 = doc.add_paragraph()
    r3h = s3.add_run("Table 3.  Residual Diagnostics")
    r3h.bold = True; r3h.font.size = Pt(12); r3h.font.name = "Arial"
    r3h.font.color.rgb = TEXT_DARK

    diag_rows = [
        ("Mean Residual",        f"{np.mean(resid_arr):.6f}"),
        ("Std Dev Residual",     f"{np.std(resid_arr):.6f}"),
        ("Skewness",             f"{skew_v:.4f}"),
        ("Excess Kurtosis",      f"{kurt_v:.4f}"),
        ("Jarque-Bera stat",     f"{jb_stat:.4f}"),
        ("Jarque-Bera p-value",  f"{jb_p:.4f}"),
        ("Durbin-Watson",        f"{dw:.4f}"),
    ]
    if bp_stat is not None:
        diag_rows += [
            ("Breusch-Pagan LM",   f"{bp_stat:.4f}"),
            ("Breusch-Pagan p",    f"{bp_p_v:.4f}"),
        ]

    diag_tbl = doc.add_table(rows=len(diag_rows) + 2, cols=2)
    clear_tbl_borders(diag_tbl)
    hdr = diag_tbl.rows[1]
    for j, h in enumerate(["Diagnostic", "Value"]):
        c = hdr.cells[j]
        set_academic_border(c, top=THICK, bottom=THIN)
        p = c.paragraphs[0]
        p.paragraph_format.space_before = p.paragraph_format.space_after = Pt(3)
        r_h = p.add_run(h)
        r_h.bold = True; r_h.font.size = Pt(9); r_h.font.name = "Arial"
        r_h.font.color.rgb = TEXT_DARK
        if j == 1:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    for i, (name, val) in enumerate(diag_rows):
        row = diag_tbl.rows[i + 2]
        bot_b = THICK if i == len(diag_rows) - 1 else NO
        for j, txt in enumerate([name, val]):
            c = row.cells[j]
            set_academic_border(c, bottom=bot_b)
            p = c.paragraphs[0]
            p.paragraph_format.space_before = p.paragraph_format.space_after = Pt(3)
            rv = p.add_run(txt)
            rv.font.size = Pt(9); rv.font.name = "Arial"
            rv.font.color.rgb = TEXT_DARK
            if j == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Section 4: AI Write-Up
    if ai_explanation.strip():
        doc.add_page_break()
        s4 = doc.add_paragraph()
        r4h = s4.add_run("AI Model Interpretation")
        r4h.bold = True; r4h.font.size = Pt(13); r4h.font.name = "Arial"
        r4h.font.color.rgb = TEAL_RGB
        doc.add_paragraph()
        for line in ai_explanation.strip().split("\n"):
            if not line.strip():
                doc.add_paragraph()
                continue
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            import re
            parts = re.split(r'\*\*(.+?)\*\*', line)
            for k, part in enumerate(parts):
                r_p = p.add_run(part)
                r_p.font.size = Pt(9.5); r_p.font.name = "Arial"
                r_p.font.color.rgb = TEXT_DARK
                if k % 2 == 1:
                    r_p.bold = True

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ═══════════════════════════════════════════════════════════════════════════════
# SIDEBAR
# ═══════════════════════════════════════════════════════════════════════════════

with st.sidebar:
    # Suite nav
    st.markdown("""
    <div style="padding:14px 0 20px 0;border-bottom:1px solid var(--border);margin-bottom:18px;">
        <div style="font-family:'Syne',sans-serif;font-size:1.3rem;font-weight:800;
                    color:var(--text);letter-spacing:-0.02em;">
            📐 Panel<span style="color:var(--accent);">Stat</span>X
        </div>
        <div style="font-family:'DM Mono',monospace;font-size:0.62rem;color:var(--muted);
                    margin-top:4px;letter-spacing:0.08em;">
            PANEL REGRESSION ENGINE · BAYANTX360 SUITE
        </div>
    </div>
    """, unsafe_allow_html=True)

    render_credit_hud()

    # Back to suite
    if st.button("⬡ Back to Suite", use_container_width=True):
        st.switch_page(st.session_state["_home_page"])

    st.markdown("---")
    st.markdown('<div style="font-size:0.68rem;text-transform:uppercase;letter-spacing:0.12em;color:var(--muted);margin-bottom:8px;">Data Source</div>', unsafe_allow_html=True)
    data_src = st.radio("", ["Use Demo Dataset", "Upload File"], label_visibility="collapsed")

    if data_src == "Upload File":
        uploaded = st.file_uploader(
            "Upload panel data", type=["csv", "xlsx", "xls"],
            label_visibility="collapsed",
        )
        if uploaded:
            try:
                fname = uploaded.name.lower()
                if fname.endswith(".csv"):
                    st.session_state.df = pd.read_csv(uploaded)
                else:
                    xf = pd.ExcelFile(uploaded)
                    sheet = st.selectbox("Sheet", xf.sheet_names) if len(xf.sheet_names) > 1 else xf.sheet_names[0]
                    st.session_state.df = pd.read_excel(uploaded, sheet_name=sheet)
                st.success(f"Loaded {st.session_state.df.shape[0]:,} rows × {st.session_state.df.shape[1]} cols")
            except Exception as e:
                st.error(f"Could not read file: {e}")
    else:
        if st.button("Load Demo Data", use_container_width=True):
            st.session_state.df = generate_demo_panel()
            st.success("Demo panel loaded!")

    st.markdown("---")

    if st.session_state.df is not None:
        df = st.session_state.df
        cols = df.columns.tolist()

        st.markdown('<div style="font-size:0.68rem;text-transform:uppercase;letter-spacing:0.12em;color:var(--muted);margin-bottom:8px;">Variable Mapping</div>', unsafe_allow_html=True)
        entity_col = st.selectbox("Entity / Panel ID", cols, index=cols.index("entity") if "entity" in cols else 0)
        time_col   = st.selectbox("Time Variable", cols, index=cols.index("year") if "year" in cols else min(1, len(cols)-1))
        y_col      = st.selectbox("Dependent Variable (Y)", [c for c in cols if c not in [entity_col, time_col]], index=0)
        x_candidates = [c for c in cols if c not in [entity_col, time_col, y_col]]
        x_cols = st.multiselect("Independent Variables (X)", x_candidates, default=x_candidates[:3] if len(x_candidates) >= 3 else x_candidates)

        st.markdown("---")
        st.markdown('<div style="font-size:0.68rem;text-transform:uppercase;letter-spacing:0.12em;color:var(--muted);margin-bottom:8px;">Estimator</div>', unsafe_allow_html=True)
        model_type = st.selectbox("", [
            "Fixed Effects (Two-Way)", "Fixed Effects (Entity)",
            "Random Effects (GLS)", "First Difference", "Pooled OLS"
        ], label_visibility="collapsed")
        st.session_state.model_type = model_type

        st.markdown("---")
        run_btn = st.button("▶ Run Analysis", use_container_width=True, type="primary")

        if st.session_state.results is not None:
            st.markdown("""
            <div style="display:flex;align-items:center;gap:8px;margin-top:8px;padding:9px 14px;
                        background:rgba(34,211,160,0.07);border:1px solid rgba(34,211,160,0.25);
                        border-radius:8px;font-family:'DM Mono',monospace;font-size:0.66rem;color:#22d3a0;">
                ✓ Analysis complete
            </div>
            """, unsafe_allow_html=True)
    else:
        run_btn = False
        entity_col = time_col = y_col = "—"
        x_cols = []
        model_type = "Fixed Effects (Two-Way)"

    st.markdown("---")
    if st.button("New Analysis", use_container_width=True):
        st.session_state.df = None
        st.session_state.results = None
        st.session_state.ai_explanation = ""
        st.rerun()

    if st.button("Sign Out", use_container_width=True):
        sign_out()


# ═══════════════════════════════════════════════════════════════════════════════
# RUN MODEL
# ═══════════════════════════════════════════════════════════════════════════════

if run_btn and st.session_state.df is not None and x_cols:
    df = st.session_state.df
    with st.spinner("Running regression…"):
        try:
            if model_type == "Pooled OLS":
                result_df, resid, y_hat, stats, vcov = run_ols(df, y_col, x_cols)
            elif model_type == "First Difference":
                result_df, resid, y_hat, stats, vcov = run_fd(df, y_col, x_cols, entity_col, time_col)
            elif model_type == "Fixed Effects (Entity)":
                result_df, resid, y_hat, stats, vcov = run_within(df, y_col, x_cols, entity_col, time_col)
            elif model_type == "Random Effects (GLS)":
                result_df, resid, y_hat, stats, vcov = run_re(df, y_col, x_cols, entity_col, time_col)
            else:
                result_df, resid, y_hat, stats, vcov = run_within(df, y_col, x_cols, entity_col, time_col)

            try:
                _bp_X = np.column_stack([np.ones(len(resid))] + [df[c].values[:len(resid)] for c in x_cols])
                bp_stat, bp_p, _ = breusch_pagan_test(resid, _bp_X)
                stats["BP_stat"] = bp_stat
                stats["BP_p"]    = bp_p
            except Exception:
                stats["BP_stat"] = None
                stats["BP_p"]    = None

            hausman_result = None
            if model_type == "Random Effects (GLS)":
                try:
                    fe_res, _, _, _, fe_vcov = run_within(df, y_col, x_cols, entity_col, time_col)
                    fe_coef = fe_res["Coeff"].values
                    re_coef = result_df[result_df["Variable"] != "const"]["Coeff"].values
                    h_stat, h_p, h_df = hausman_test(fe_coef, re_coef, fe_vcov, vcov[1:, 1:])
                    hausman_result = {"stat": h_stat, "p": h_p, "df": h_df}
                except Exception:
                    pass

            st.session_state.results = {
                "result_df": result_df, "resid": resid, "y_hat": y_hat,
                "stats": stats, "y_col": y_col, "x_cols": x_cols,
                "entity_col": entity_col, "time_col": time_col,
                "hausman": hausman_result,
            }
            st.session_state.ai_explanation = ""
            st.rerun()

        except Exception as e:
            st.error(f"Regression error: {e}")


# ═══════════════════════════════════════════════════════════════════════════════
# MAIN LAYOUT
# ═══════════════════════════════════════════════════════════════════════════════

st.markdown("""
<div class="app-hero">
    <div class="app-hero-title">Panel<span>Stat</span>X</div>
    <div class="app-hero-sub">📐 Panel Regression Analysis System · Bayantx360 Suite</div>
</div>
""", unsafe_allow_html=True)

if st.session_state.df is None:
    c1, c2, c3 = st.columns(3)
    for col, icon, title, desc in [
        (c1, "⬡", "Panel-Ready", "Fixed Effects, Random Effects, First-Difference, and Pooled OLS estimators built for longitudinal data."),
        (c2, "◈", "Diagnostic Suite", "Residual analysis, Breusch-Pagan heteroskedasticity, Hausman specification test, and entity plots."),
        (c3, "⬟", "AI Explainer", "AI model interprets your regression output in plain language — coefficients, fit, and caveats."),
    ]:
        with col:
            st.markdown(f"""
            <div class="scard" style="text-align:center;padding:32px 20px;">
                <div style="font-size:2rem;margin-bottom:12px;color:var(--accent);">{icon}</div>
                <div style="font-family:'Syne',sans-serif;font-size:1rem;font-weight:700;color:var(--text);margin-bottom:8px;">{title}</div>
                <div style="font-size:0.76rem;color:var(--muted);line-height:1.7;">{desc}</div>
            </div>
            """, unsafe_allow_html=True)
    st.markdown("""
    <div style="text-align:center;margin-top:40px;padding:32px;background:var(--surface2);
                border:1px dashed var(--border);border-radius:10px;">
        <div style="font-family:'Syne',sans-serif;font-size:1rem;color:var(--muted);">
            ← Load demo data or upload a CSV from the sidebar to begin
        </div>
    </div>
    """, unsafe_allow_html=True)
    st.stop()

df  = st.session_state.df
res = st.session_state.results

n_e = df[entity_col].nunique() if entity_col in df.columns else "—"
n_t = df[time_col].nunique()   if time_col   in df.columns else "—"
st.markdown(f"""
<div style="margin-bottom:22px;">
    <span class="stat-pill">Entities <b>{n_e}</b></span>
    <span class="stat-pill">Periods <b>{n_t}</b></span>
    <span class="stat-pill">Observations <b>{len(df):,}</b></span>
    <span class="stat-pill">Estimator <b>{st.session_state.model_type}</b></span>
    <span class="badge badge-teal" style="margin-left:4px;">READY</span>
    {"<span class='badge badge-warn' style='margin-left:4px;'>FREE TRIAL</span>" if is_trial() else ""}
</div>
""", unsafe_allow_html=True)

tab_data, tab_results, tab_diagnostics, tab_entity, tab_ai = st.tabs([
    "⬡ Data Explorer", "◈ Results", "⬟ Diagnostics", "⬢ Entity Plots", "✦ AI Explainer"
])

# ─────────────────── TAB 1: DATA EXPLORER ─────────────────────────────────────
with tab_data:
    c1, c2 = st.columns([2, 1])
    with c1:
        st.markdown('<div class="scard-title">Dataset Preview</div>', unsafe_allow_html=True)
        st.dataframe(df.head(100), use_container_width=True, height=320)
    with c2:
        st.markdown('<div class="scard-title">Summary Statistics</div>', unsafe_allow_html=True)
        st.dataframe(df.describe().round(3), use_container_width=True, height=320)

    st.markdown("---")
    num_cols = df.select_dtypes(include=[np.number]).columns.tolist()
    if len(num_cols) >= 2:
        st.markdown('<div class="scard-title">Correlation Heatmap</div>', unsafe_allow_html=True)
        corr = df[num_cols].corr().round(3)
        fig_corr = go.Figure(go.Heatmap(
            z=corr.values, x=corr.columns, y=corr.index,
            colorscale=[[0, "#f05c7c"], [0.5, "#111318"], [1, "#00e5c8"]],
            zmid=0, text=corr.values.round(2), texttemplate="%{text}", showscale=True,
        ))
        apply_theme(fig_corr.update_layout(title="Pearson Correlation Matrix", height=380))
        st.plotly_chart(fig_corr, use_container_width=True)

    if y_col in df.columns:
        st.markdown('<div class="scard-title" style="margin-top:16px;">Dependent Variable Distribution</div>', unsafe_allow_html=True)
        fig_dist = px.histogram(df, x=y_col, nbins=40, color_discrete_sequence=["#00e5c8"])
        apply_theme(fig_dist.update_layout(title=f"Distribution of {y_col}", height=300, bargap=0.05))
        st.plotly_chart(fig_dist, use_container_width=True)


# ─────────────────── TAB 2: RESULTS ───────────────────────────────────────────
with tab_results:
    if res is None:
        st.info("Run the analysis from the sidebar to view regression results.")
    else:
        result_df = res["result_df"]
        stats = res["stats"]

        st.markdown('<div class="scard-title">Model Fit</div>', unsafe_allow_html=True)
        mc = st.columns(8)
        f_label  = f"{stats.get('F_stat', 0):.3f}" if stats.get('F_stat') is not None else "—"
        fp_label = f"{stats.get('F_p', 1):.4f}"    if stats.get('F_p')    is not None else "—"
        for col, label, val in [
            (mc[0], "R²",        f"{stats['R2']:.4f}"),
            (mc[1], "Adj. R²",   f"{stats['R2_adj']:.4f}"),
            (mc[2], "N",         f"{stats['N']:,}"),
            (mc[3], "Variables", f"{stats['k']}"),
            (mc[4], "AIC",       f"{stats['AIC']:.2f}"),
            (mc[5], "BIC",       f"{stats['BIC']:.2f}"),
            (mc[6], "F-stat",    f_label),
            (mc[7], "F p-value", fp_label),
        ]:
            with col: st.metric(label, val)

        if stats.get("F_p") is not None:
            if stats["F_p"] < 0.05:
                st.success(f"✓ F-statistic ({stats['F_stat']:.3f}) is significant (p={stats['F_p']:.4f})")
            else:
                st.warning(f"⚠ F-statistic not significant (p={stats['F_p']:.4f})")

        if st.session_state.model_type == "Random Effects (GLS)":
            st.markdown("---")
            st.markdown('<div class="scard-title">Random Effects Variance Components</div>', unsafe_allow_html=True)
            rc1, rc2, rc3 = st.columns(3)
            with rc1: st.metric("σ²ᵤ (between)", f"{stats.get('sigma_u2', 0):.6f}")
            with rc2: st.metric("σ²ₑ (within)",  f"{stats.get('sigma_e2', 0):.6f}")
            with rc3: st.metric("θ (GLS weight)", f"{stats.get('theta', 0):.4f}")
            st.caption("θ → 1: FE dominates. θ → 0: OLS/pooled dominates.")

        hausman_res = res.get("hausman")
        if hausman_res and hausman_res.get("stat") is not None:
            st.markdown("---")
            st.markdown('<div class="scard-title">Hausman Specification Test (RE vs FE)</div>', unsafe_allow_html=True)
            hc1, hc2, hc3 = st.columns(3)
            with hc1: st.metric("χ² statistic", f"{hausman_res['stat']:.4f}")
            with hc2: st.metric("p-value",       f"{hausman_res['p']:.4f}")
            with hc3: st.metric("df",             f"{hausman_res['df']}")
            if hausman_res["p"] < 0.05:
                st.warning("⚠ Hausman rejects RE (p < 0.05) — Fixed Effects preferred.")
            else:
                st.success("✓ Hausman does not reject RE (p ≥ 0.05) — RE consistent and efficient.")

        st.markdown("---")
        st.markdown('<div class="scard-title">Coefficient Estimates</div>', unsafe_allow_html=True)
        display = result_df.copy()
        display["Stars"] = display["p_value"].apply(significance_stars)
        display["Sig"]   = display["p_value"].apply(lambda p: "✓ Significant" if p < 0.05 else "✗ Not sig.")
        display = display.rename(columns={
            "Variable": "Variable", "Coeff": "Coeff.", "Std_Err": "Std. Err.",
            "t_stat": "t-stat", "p_value": "p-value"
        })
        st.dataframe(
            display.style
                .format({"Coeff.": "{:.4f}", "Std. Err.": "{:.4f}", "t-stat": "{:.3f}", "p-value": "{:.4f}"})
                .map(lambda v: "color: #00e5c8" if v == "✓ Significant" else "color: #6b7a9a", subset=["Sig"]),
            use_container_width=True, hide_index=True
        )
        st.caption("*p<0.1  **p<0.05  ***p<0.01")

        st.markdown("---")
        st.markdown('<div class="scard-title">Coefficient Plot (95% CI)</div>', unsafe_allow_html=True)
        rd = res["result_df"]
        ci_lo = rd["Coeff"] - 1.96 * rd["Std_Err"]
        ci_hi = rd["Coeff"] + 1.96 * rd["Std_Err"]
        colors = ["#00e5c8" if p < 0.05 else "#6b7a9a" for p in rd["p_value"]]
        fig_coef = go.Figure()
        fig_coef.add_hline(y=0, line_dash="dash", line_color="#1f2535")
        for pos, (i, row_r) in enumerate(rd.iterrows()):
            fig_coef.add_trace(go.Scatter(
                x=[ci_lo.iloc[pos], ci_hi.iloc[pos]], y=[row_r["Variable"], row_r["Variable"]],
                mode="lines", line=dict(color="#1f2535", width=2), showlegend=False
            ))
        fig_coef.add_trace(go.Scatter(
            x=rd["Coeff"], y=rd["Variable"], mode="markers",
            marker=dict(size=10, color=colors, line=dict(width=1, color="#0a0c10")),
            showlegend=False
        ))
        apply_theme(fig_coef.update_layout(title="Coefficients with 95% Confidence Intervals",
                                            height=max(280, len(rd) * 55)))
        st.plotly_chart(fig_coef, use_container_width=True)

        st.markdown("---")
        st.markdown('<div class="scard-title">Export Report</div>', unsafe_allow_html=True)

        if is_trial():
            render_locked_banner("DOCX Report Export", is_trial_user=True)
            _, uc, _ = st.columns([1, 2, 1])
            with uc:
                st.link_button("Upgrade — Get Access Key →", "https://x.com/bayantx360", use_container_width=True)
        else:
            dl_col1, dl_col2 = st.columns([1, 2])
            with dl_col1:
                credits_left = st.session_state.user_credits
                dl_disabled  = credits_left <= 0
                try:
                    import datetime
                    docx_bytes = build_docx_report(
                        res, st.session_state.model_type,
                        ai_explanation=st.session_state.get("ai_explanation", ""),
                    )
                    fname = f"PanelStatX_Report_{datetime.datetime.now().strftime('%Y%m%d_%H%M')}.docx"
                    if st.download_button(
                        label="⬇  Download Report (.docx)",
                        data=docx_bytes, file_name=fname,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True, type="primary", disabled=dl_disabled,
                    ):
                        handle_credit_deduction(1, app="PanelStatX", action="DOCX Report Export")
                        st.rerun()
                    if dl_disabled:
                        st.caption("⚠ No credits remaining.")
                    else:
                        st.caption(f"⚡ Costs 1 credit · {credits_left} remaining")
                except Exception as e:
                    st.error(f"Report generation error: {e}")

            with dl_col2:
                ai_ready = bool(st.session_state.get("ai_explanation", "").strip())
                tip_color = "#22d3a0" if ai_ready else "#6b7a9a"
                tip_text  = "✓ AI interpretation included in report." if ai_ready else "○ Generate AI explanation first, then download."
                st.markdown(f'<div style="font-family:\'DM Mono\',monospace;font-size:0.7rem;color:{tip_color};padding:10px 0;line-height:1.8;">{tip_text}</div>', unsafe_allow_html=True)


# ─────────────────── TAB 3: DIAGNOSTICS ───────────────────────────────────────
with tab_diagnostics:
    if res is None:
        st.info("Run the analysis first.")
    else:
        resid = res["resid"]
        y_hat = res["y_hat"]
        dc1, dc2 = st.columns(2)
        with dc1:
            fig_rv = go.Figure()
            fig_rv.add_hline(y=0, line_dash="dash", line_color="#f05c7c", line_width=1)
            fig_rv.add_trace(go.Scatter(x=y_hat, y=resid, mode="markers",
                                         marker=dict(size=5, color="#00e5c8", opacity=0.6)))
            apply_theme(fig_rv.update_layout(title="Residuals vs Fitted",
                                             xaxis_title="Fitted", yaxis_title="Residual", height=340))
            st.plotly_chart(fig_rv, use_container_width=True)

        with dc2:
            from scipy import stats as sc_stats
            resid_clean = np.asarray(resid, dtype=float)
            resid_clean = resid_clean[np.isfinite(resid_clean)]
            pp = sc_stats.probplot(resid_clean)
            osm, osr = pp[0]
            slope, intercept, _ = pp[1]
            fig_qq = go.Figure()
            fig_qq.add_trace(go.Scatter(x=list(osm), y=list(osr), mode="markers",
                                         marker=dict(size=4, color="#7c6df0", opacity=0.7)))
            fig_qq.add_trace(go.Scatter(
                x=[float(min(osm)), float(max(osm))],
                y=[slope * float(min(osm)) + intercept, slope * float(max(osm)) + intercept],
                mode="lines", line=dict(color="#f05c7c", dash="dash")
            ))
            apply_theme(fig_qq.update_layout(title="Normal Q-Q Plot",
                                             xaxis_title="Theoretical Quantiles",
                                             yaxis_title="Sample Quantiles", height=340))
            st.plotly_chart(fig_qq, use_container_width=True)

        dc3, dc4 = st.columns(2)
        with dc3:
            fig_rh = px.histogram(x=resid, nbins=40, color_discrete_sequence=["#7c6df0"])
            apply_theme(fig_rh.update_layout(title="Residual Distribution", height=300, bargap=0.05))
            st.plotly_chart(fig_rh, use_container_width=True)
        with dc4:
            fig_sl = go.Figure()
            fig_sl.add_trace(go.Scatter(x=y_hat, y=np.sqrt(np.abs(resid)), mode="markers",
                                         marker=dict(size=5, color="#f5a623", opacity=0.6)))
            apply_theme(fig_sl.update_layout(title="Scale-Location", height=300))
            st.plotly_chart(fig_sl, use_container_width=True)

        st.markdown("---")
        st.markdown('<div class="scard-title">Residual Diagnostics Summary</div>', unsafe_allow_html=True)
        from scipy import stats as sc_stats
        resid_arr = np.asarray(resid, dtype=float)
        resid_arr = resid_arr[np.isfinite(resid_arr)]
        jb_stat, jb_p = sc_stats.jarque_bera(resid_arr)
        dw = np.sum(np.diff(resid_arr)**2) / max(np.sum(resid_arr**2), 1e-10)
        dc5, dc6, dc7, dc8 = st.columns(4)
        with dc5: st.metric("Mean Residual", f"{np.mean(resid_arr):.4f}")
        with dc6: st.metric("Std Residual",  f"{np.std(resid_arr):.4f}")
        with dc7: st.metric("Jarque-Bera p", f"{jb_p:.4f}")
        with dc8: st.metric("Durbin-Watson", f"{dw:.4f}")

        if jb_p < 0.05:
            st.warning("⚠ Jarque-Bera rejects normality (p < 0.05).")
        if dw < 1.5 or dw > 2.5:
            st.warning(f"⚠ Durbin-Watson = {dw:.3f} — possible autocorrelation.")
        else:
            st.success("✓ Durbin-Watson in acceptable range (1.5–2.5).")

        st.markdown("---")
        st.markdown('<div class="scard-title">Breusch-Pagan Test for Heteroskedasticity</div>', unsafe_allow_html=True)
        bp_stat = res["stats"].get("BP_stat")
        bp_p    = res["stats"].get("BP_p")
        if bp_stat is not None:
            bpc1, bpc2 = st.columns(2)
            with bpc1: st.metric("BP LM Statistic", f"{bp_stat:.4f}")
            with bpc2: st.metric("BP p-value",       f"{bp_p:.4f}")
            if bp_p < 0.05:
                st.warning("⚠ Heteroskedastic errors detected — consider robust standard errors.")
            else:
                st.success("✓ Homoskedasticity not rejected (p ≥ 0.05).")
            st.caption("H₀: Constant variance. LM ~ χ²(k).")
        else:
            st.info("Breusch-Pagan test unavailable for this estimator.")


# ─────────────────── TAB 4: ENTITY PLOTS ──────────────────────────────────────
with tab_entity:
    if entity_col not in df.columns or time_col not in df.columns:
        st.info("Entity and time columns not set.")
    else:
        y_plot = y_col if y_col in df.columns else df.select_dtypes(np.number).columns[0]
        ec1, ec2 = st.columns([1, 3])
        with ec1:
            entities_avail = sorted(df[entity_col].unique())
            selected_entities = st.multiselect("Select entities", entities_avail,
                                               default=entities_avail[:6] if len(entities_avail) >= 6 else entities_avail)
        with ec2:
            x_axis = st.selectbox("X axis", [time_col] + [c for c in df.columns if c not in [entity_col]], index=0)

        if selected_entities:
            plot_df = df[df[entity_col].isin(selected_entities)]
            fig_ep = px.line(plot_df, x=x_axis, y=y_plot, color=entity_col, markers=True,
                             color_discrete_sequence=["#00e5c8","#7c6df0","#f05c7c","#f5a623","#22d3a0","#60a5fa","#fb923c","#a78bfa"])
            apply_theme(fig_ep.update_layout(title=f"{y_plot} over {x_axis}", height=440))
            st.plotly_chart(fig_ep, use_container_width=True)

            means = df.groupby(entity_col)[y_plot].mean().sort_values(ascending=False)
            fig_bar = px.bar(x=means.index, y=means.values,
                             color=means.values, color_continuous_scale=["#111318","#00e5c8"],
                             labels={"x": entity_col, "y": f"Mean {y_plot}"})
            apply_theme(fig_bar.update_layout(title=f"Entity Mean of {y_plot}", height=320, coloraxis_showscale=False))
            st.plotly_chart(fig_bar, use_container_width=True)


# ─────────────────── TAB 5: AI EXPLAINER ──────────────────────────────────────
with tab_ai:
    st.markdown('<div class="scard-title">AI Regression Explainer</div>', unsafe_allow_html=True)

    if res is None:
        st.info("Run the analysis first to unlock the AI explainer.")
    else:
        result_df = res["result_df"]
        stats = res["stats"]
        coeff_table = result_df.to_string(index=False)
        hausman_res = res.get("hausman")
        hausman_str = ""
        if hausman_res and hausman_res.get("stat") is not None:
            hausman_str = f"\nHausman Test: χ²={hausman_res['stat']:.4f}, p={hausman_res['p']:.4f} (df={hausman_res['df']})"
        from scipy import stats as _sc_stats
        _resid_arr = np.asarray(res["resid"], dtype=float)
        _resid_arr = _resid_arr[np.isfinite(_resid_arr)]
        _jb_stat, _jb_p = _sc_stats.jarque_bera(_resid_arr)
        _dw = np.sum(np.diff(_resid_arr)**2) / max(np.sum(_resid_arr**2), 1e-10)
        _skew = _sc_stats.skew(_resid_arr)
        _kurt = _sc_stats.kurtosis(_resid_arr)
        _bp_s = stats.get("BP_stat")
        _bp_p_v = stats.get("BP_p")
        _bp_str = f"LM={_bp_s:.4f}, p={_bp_p_v:.4f}" if _bp_s is not None else "N/A"
        re_str = ""
        if st.session_state.model_type == "Random Effects (GLS)":
            re_str = f"\nσ²ᵤ={stats.get('sigma_u2',0):.6f}, σ²ₑ={stats.get('sigma_e2',0):.6f}, θ={stats.get('theta',0):.4f}"

        context = f"""
Model: {st.session_state.model_type}
Dependent: {res['y_col']} | Independent: {', '.join(res['x_cols'])}
Entity: {res['entity_col']} | Time: {res['time_col']}
R²={stats['R2']:.4f}, Adj.R²={stats['R2_adj']:.4f}, N={stats['N']}, AIC={stats['AIC']:.2f}, BIC={stats['BIC']:.2f}
F={stats.get('F_stat',float('nan')):.4f} (p={stats.get('F_p',float('nan')):.4f}){hausman_str}{re_str}
Jarque-Bera={_jb_stat:.4f} (p={_jb_p:.4f}), DW={_dw:.4f}, Skew={_skew:.4f}, Kurt={_kurt:.4f}
Breusch-Pagan={_bp_str}
Coefficients:
{coeff_table}"""

        sys_prompt = (
            "You are an expert econometrician. Produce a structured, publication-quality interpretation. "
            "Use EXACT numeric values from the input. Format with these sections:\n"
            "## Model Specification\n## Coefficient Interpretation\n"
            "## Model Fit & Overall Significance\n## Diagnostic Test Results\n"
            "## Caveats & Concerns\n## Recommendations\n"
            "Use **bold** for key terms and numbers. No preamble or closing remarks."
        )

        if is_trial():
            render_locked_banner("AI Explainer", is_trial_user=True)
            _, uc, _ = st.columns([1, 2, 1])
            with uc:
                st.link_button("⬡  Upgrade to Paid Plan →", "https://x.com/bayantx360", use_container_width=True)
        else:
            credits_left = st.session_state.user_credits
            col_explain, col_custom = st.columns([3, 2])
            with col_explain:
                ai_disabled = credits_left <= 0
                if st.button("✦ Generate AI Explanation", type="primary", use_container_width=True, disabled=ai_disabled):
                    with st.spinner("AI model analysing results…"):
                        explanation = call_openai(sys_prompt, f"Explain these results:\n\n{context}")
                        if not explanation.startswith(("OpenAI API key", "Request failed")):
                            handle_credit_deduction(1, app="PanelStatX", action="AI Explainer")
                        st.session_state.ai_explanation = explanation
                        st.rerun()
                if ai_disabled:
                    st.caption("⚠ No credits remaining.")
                else:
                    st.caption(f"⚡ Costs 1 credit · {credits_left} remaining")

            with col_custom:
                custom_q = st.text_input("Ask a specific question…", placeholder="e.g. Is x1 economically significant?")
                ask_disabled = credits_left <= 0
                if st.button("Ask AI", use_container_width=True, disabled=ask_disabled) and custom_q:
                    with st.spinner("Thinking…"):
                        answer = call_openai(sys_prompt, f"Results:\n\n{context}\n\nQuestion: {custom_q}")
                        if not answer.startswith(("OpenAI API key", "Request failed")):
                            handle_credit_deduction(1, app="PanelStatX", action="AI Custom Question")
                        st.session_state.ai_explanation = answer
                        st.rerun()
                if ask_disabled:
                    st.caption("⚠ No credits remaining.")
                else:
                    st.caption(f"⚡ Costs 1 credit · {credits_left} remaining")

            if st.session_state.ai_explanation:
                st.markdown("---")
                import re as _re, html as _html

                def _md_to_html(text):
                    HDR  = "font-family:'Syne',sans-serif;font-size:0.78rem;font-weight:700;text-transform:uppercase;letter-spacing:0.1em;color:var(--accent);margin:18px 0 6px;border-bottom:1px solid rgba(0,229,200,0.2);padding-bottom:4px;"
                    out  = []
                    for line in text.split("\n"):
                        s = line.strip()
                        if not s: out.append("<br>"); continue
                        if s.startswith("## ") or s.startswith("### "):
                            out.append(f'<div style="{HDR}">{_html.escape(s.lstrip("# ").rstrip(":"))}</div>'); continue
                        s_e = _html.escape(s)
                        s_e = _re.sub(r'\*\*([^*]+)\*\*', r'<strong>\1</strong>', s_e)
                        s_e = _re.sub(r'\*([^*]+)\*', r'<em>\1</em>', s_e)
                        if _re.match(r'^[-\u2022]\s+', s_e):
                            body = _re.sub(r'^[-\u2022]\s+', '', s_e)
                            out.append(f'<div style="padding:2px 0 2px 18px;position:relative;"><span style="position:absolute;left:4px;color:var(--accent);">–</span>{body}</div>'); continue
                        m = _re.match(r'^(\d+)[.)]\s+(.*)', s_e)
                        if m:
                            out.append(f'<div style="padding:2px 0 2px 24px;position:relative;"><span style="position:absolute;left:4px;color:var(--accent);font-weight:700;">{m.group(1)}.</span>{m.group(2)}</div>'); continue
                        out.append(f'<div style="padding:2px 0;">{s_e}</div>')
                    return "\n".join(out)

                ai_html = _md_to_html(st.session_state.ai_explanation)
                st.markdown(
                    '<div class="ai-label">✦ &nbsp;AI MODEL INTERPRETATION</div>'
                    '<div class="ai-box" style="white-space:normal;">' + ai_html + '</div>',
                    unsafe_allow_html=True,
                )

        st.markdown("---")
        st.markdown('<div class="scard-title">Quick Insights</div>', unsafe_allow_html=True)
        ic1, ic2, ic3 = st.columns(3)
        sig_vars   = result_df[result_df["p_value"] < 0.05]["Variable"].tolist()
        r2_val     = stats["R2"]
        r2_color   = "#00e5c8" if r2_val > 0.7 else "#f5a623" if r2_val > 0.4 else "#f05c7c"
        r2_label   = "Strong fit" if r2_val > 0.7 else "Moderate fit" if r2_val > 0.4 else "Weak fit"
        largest    = result_df.iloc[result_df["Coeff"].abs().argmax()]
        with ic1:
            st.markdown(f'<div class="scard"><div class="scard-title">Significance</div><div style="color:var(--accent);font-size:1.2rem;font-family:\'Syne\',sans-serif;font-weight:700;">{len(sig_vars)}/{len(result_df)}</div><div style="color:var(--muted);font-size:0.72rem;margin-top:4px;">significant at 5%</div><div style="margin-top:8px;font-size:0.7rem;color:var(--text);">{", ".join(sig_vars) if sig_vars else "—"}</div></div>', unsafe_allow_html=True)
        with ic2:
            st.markdown(f'<div class="scard"><div class="scard-title">Model Fit</div><div style="color:{r2_color};font-size:1.2rem;font-family:\'Syne\',sans-serif;font-weight:700;">{r2_val:.4f}</div><div style="color:var(--muted);font-size:0.72rem;margin-top:4px;">R² · {r2_label}</div></div>', unsafe_allow_html=True)
        with ic3:
            st.markdown(f'<div class="scard"><div class="scard-title">Largest Effect</div><div style="color:var(--accent2);font-size:1.2rem;font-family:\'Syne\',sans-serif;font-weight:700;">{largest["Variable"]}</div><div style="color:var(--muted);font-size:0.72rem;margin-top:4px;">coeff = {largest["Coeff"]:.4f}</div></div>', unsafe_allow_html=True)

st.markdown("---")
st.markdown('<div style="text-align:center;font-family:\'DM Mono\',monospace;font-size:0.68rem;color:var(--muted);padding:10px 0;">📐 PanelStatX · Panel Regression Analysis System · Bayantx360 Suite</div>', unsafe_allow_html=True)
