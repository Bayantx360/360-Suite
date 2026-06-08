"""
apps/efactor.py
══════════════════════════════════════════════════════════════════════
Bayantx360 Suite — EFActor
Psychometric Analysis Platform
══════════════════════════════════════════════════════════════════════

Changes from standalone version (app__26_.py):
  • Auth + credit engine → shared/auth.py
  • CSS → shared/theme.py (unified suite visual system)
  • Single secret: BAYANTX_SHEET_ID
  • Trial gate: is_free_trial flag (standardised)
  • Credit deduction → handle_credit_deduction() from shared.auth
  • Back-to-suite navigation added in sidebar
  • export_credit_cost() preserved (row-based tiered pricing)

All EFA/CFA logic (factor_analyzer, _ensure_psd_dataframe, auto-fix
engine, CFA, semopy, synthetic generation, DOCX builder) 100%
preserved from original app__26_.py.
"""

import io
import zipfile
import warnings
from datetime import datetime

import numpy as np
import pandas as pd
import plotly.graph_objects as go
from plotly.subplots import make_subplots
import streamlit as st
from factor_analyzer import FactorAnalyzer, calculate_kmo, calculate_bartlett_sphericity
import sys, os
warnings.filterwarnings("ignore")

# ── sklearn >= 1.6 compatibility ──────────────────────────────────────────────
try:
    import factor_analyzer.factor_analyzer as _fa_mod
    import factor_analyzer.confirmatory_factor_analyzer as _cfa_mod
    from sklearn.utils.validation import check_array as _orig_check_array
    def _compat_check_array(*args, **kwargs):
        if "force_all_finite" in kwargs:
            val = kwargs.pop("force_all_finite")
            kwargs["ensure_all_finite"] = (val is True)
        return _orig_check_array(*args, **kwargs)
    _fa_mod.check_array  = _compat_check_array
    _cfa_mod.check_array = _compat_check_array
except Exception:
    pass


sys.path.insert(0, os.path.join(os.path.dirname(os.path.abspath(__file__)), ".."))
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from shared.auth import (
    init_session_state,
    refresh_credits,
    handle_credit_deduction,
    render_credit_hud,
    is_trial,
    sign_out,
)
from shared.theme import apply_suite_css, render_locked_banner

# ── Page config ────────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="EFActor · Bayantx360 Suite",
    page_icon="🔬",
    layout="wide",
    initial_sidebar_state="expanded",
)

init_session_state()
apply_suite_css()

if not st.session_state.get("access_granted"):
    st.switch_page(st.session_state["_home_page"])
    st.stop()

refresh_credits()

# ── EFActor colour constants (preserved) ──────────────────────────────────────
C = dict(
    accent="#00e5c8", accent2="#7c6df0",
    green="#22d3a0", red="#f05c7c", yellow="#f5a623",
    bg="#0a0c10", surface="#111318", surface2="#181c24", border="#1f2535",
    text="#e2e8f4", muted="#6b7a9a",
)
LAYOUT_BASE = dict(
    paper_bgcolor=C["surface"], plot_bgcolor=C["surface"],
    font=dict(color=C["text"], family="DM Mono, monospace"),
    margin=dict(l=40, r=20, t=50, b=40),
    colorway=[C["accent"], C["accent2"], C["green"], C["yellow"], C["red"]],
)

# ── Per-app session keys ───────────────────────────────────────────────────────
_DEFAULTS = dict(
    df_original=None, df_working=None, suitability=None,
    n_factors_auto=None, eigenvalues=None, efa_result=None,
    diagnostics=None, efa_done=False, cfa_result=None,
    fit_assessment=None, synthetic_factor=None, synthetic_corr=None,
    syn_validation=None, report_docx=None,
    df_autofix=None, fix_log=None, autofix_efa_result=None,
)
for _k, _v in _DEFAULTS.items():
    if _k not in st.session_state:
        st.session_state[_k] = _v
if "dropped_vars" not in st.session_state:
    st.session_state["dropped_vars"] = []
S = st.session_state


# ═══════════════════════════════════════════════════════════════════════════════
# CORE FUNCTIONS (100% preserved from original)
# ═══════════════════════════════════════════════════════════════════════════════

def export_credit_cost(n_rows: int) -> int:
    if n_rows <= 300:  return 1
    if n_rows <= 1000: return 2
    return 5

def _ensure_psd_dataframe(df: pd.DataFrame, eps: float = 1e-4) -> pd.DataFrame:
    df = df.copy()
    df = df.loc[:, df.std() > 1e-9]
    if df.shape[1] < 2: return df
    corr = df.corr().values
    eigvals = np.linalg.eigvalsh(corr)
    if eigvals.min() > eps: return df
    ev, evec = np.linalg.eigh(corr)
    ev_fixed = np.maximum(ev, eps)
    corr_fixed = evec @ np.diag(ev_fixed) @ evec.T
    d = np.sqrt(np.diag(corr_fixed))
    corr_fixed = corr_fixed / np.outer(d, d)
    np.fill_diagonal(corr_fixed, 1.0)
    try:
        L = np.linalg.cholesky(corr_fixed + np.eye(len(corr_fixed)) * eps)
    except np.linalg.LinAlgError:
        ridge = eps
        for _ in range(20):
            try:
                L = np.linalg.cholesky(corr_fixed + np.eye(len(corr_fixed)) * ridge); break
            except np.linalg.LinAlgError:
                ridge *= 10
        else:
            return df
    rng = np.random.default_rng(0)
    z = rng.standard_normal((len(df), len(df.columns)))
    z_corr = z @ L.T
    stds, means = df.std().values, df.mean().values
    z_norm = (z_corr - z_corr.mean(axis=0)) / np.maximum(z_corr.std(axis=0), 1e-9)
    return pd.DataFrame(z_norm * stds + means, columns=df.columns, index=df.index)

def check_efa_suitability(df):
    kmo_all, kmo_model = calculate_kmo(df)
    chi2, p = calculate_bartlett_sphericity(df)
    labels = {.9:"Marvellous",.8:"Meritorious",.7:"Middling",.6:"Mediocre",.5:"Miserable"}
    kmo_label = next((v for k, v in sorted(labels.items(), reverse=True) if kmo_model >= k), "Unacceptable")
    return dict(kmo_model=round(float(kmo_model),4), kmo_label=kmo_label, kmo_pass=kmo_model>=0.6,
                bartlett_chi2=round(float(chi2),4), bartlett_p=round(float(p),6),
                bartlett_pass=p<0.05, overall_pass=kmo_model>=0.6 and p<0.05)

def determine_n_factors(df):
    df = _ensure_psd_dataframe(df)
    max_factors = max(1, min(len(df.columns)-1, len(df)-1))
    try:
        fa = FactorAnalyzer(n_factors=max_factors, rotation=None); fa.fit(df)
        ev, _ = fa.get_eigenvalues()
    except Exception:
        corr = df.corr().values; ev = np.sort(np.linalg.eigvalsh(corr))[::-1]
    return dict(eigenvalues=ev.tolist(), suggested_n=max(1, int(np.sum(np.array(ev) > 1))))

def run_efa(df, n_factors, rotation="varimax"):
    df = _ensure_psd_dataframe(df)
    n_factors = min(n_factors, len(df.columns)-1)
    try:
        fa = FactorAnalyzer(n_factors=n_factors, rotation=rotation); fa.fit(df)
        factor_labels = [f"F{i+1}" for i in range(n_factors)]
        loadings = pd.DataFrame(fa.loadings_, index=df.columns, columns=factor_labels)
        communalities = pd.Series(fa.get_communalities(), index=df.columns, name="Communality")
        variance = pd.DataFrame(fa.get_factor_variance(),
                                index=["SS Loadings","Proportion Var","Cumulative Var"],
                                columns=factor_labels).T
    except Exception:
        corr = df.corr().values; ev, evec = np.linalg.eigh(corr)
        idx = np.argsort(ev)[::-1][:n_factors]; ev_top = np.maximum(ev[idx], 0)
        raw_loadings = evec[:, idx] * np.sqrt(ev_top)
        factor_labels = [f"F{i+1}" for i in range(n_factors)]
        loadings = pd.DataFrame(raw_loadings, index=df.columns, columns=factor_labels)
        communalities = pd.Series(np.clip((raw_loadings**2).sum(axis=1),0,1), index=df.columns, name="Communality")
        variance = pd.DataFrame(np.zeros((n_factors,3)), index=factor_labels,
                                columns=["SS Loadings","Proportion Var","Cumulative Var"])
    return dict(loadings=loadings, communalities=communalities, variance=variance, n_factors=n_factors)

def diagnose_loadings(loadings, communalities, load_thresh=0.4, comm_thresh=0.3):
    records = []
    for var in loadings.index:
        abs_row = np.abs(loadings.loc[var]); max_load = abs_row.max()
        n_high = int((abs_row >= load_thresh).sum()); comm = float(communalities[var])
        issues, severity = [], 0.0
        if comm < comm_thresh: issues.append("Low Communality"); severity += (comm_thresh - comm)*3
        if n_high == 0: issues.append("Weak Loader"); severity += (load_thresh - max_load)*2
        elif n_high > 1:
            issues.append("Cross-Loader"); s = sorted(abs_row.values, reverse=True)
            severity += (1 - (s[0] - s[1]))*2
        records.append(dict(Variable=var, MaxLoading=round(max_load,4),
                            FactorsAboveThreshold=n_high, Communality=round(comm,4),
                            Issue=", ".join(issues) if issues else "OK",
                            Severity=round(severity,4), RecommendDrop=len(issues)>0))
    return pd.DataFrame(records).sort_values("Severity", ascending=False).reset_index(drop=True)

# ── Auto-Fix helpers (preserved) ──────────────────────────────────────────────
def _winsorize(s, limits=(0.05,0.05)):
    return s.clip(lower=s.quantile(limits[0]), upper=s.quantile(1-limits[1]))
def _log_transform(s): return np.log1p(s - s.min() + 1e-6)
def _add_jitter(s, seed=42, scale=0.01):
    rng = np.random.default_rng(seed)
    return s + rng.normal(0, max(s.std()*scale, 1e-6), len(s))
def _rescale_to_original(fixed, original):
    om, os_ = float(original.mean()), float(original.std())
    fs = float(fixed.std())
    if fs > 1e-9 and os_ > 1e-9:
        fixed = (fixed - fixed.mean()) / fs * os_ + om
    return fixed

def _detect_data_issues(df):
    issues = {}; corr = df.corr().abs()
    for var in df.columns:
        s = df[var]; vi = []
        if s.std() < 1e-6: vi.append("zero_variance"); issues[var] = vi; continue
        z = (s - s.mean()) / s.std(); n_out = int((z.abs() > 3).sum())
        if n_out > 0: vi.append(f"outliers:{n_out}")
        skw = float(s.skew())
        if abs(skw) > 2: vi.append(f"skewness:{skw:.3f}")
        krt = float(s.kurt())
        if krt > 7: vi.append(f"kurtosis:{krt:.3f}")
        others = corr[var].drop(var)
        if len(others) and others.max() >= 0.95:
            vi.append(f"collinear:{others.idxmax()}:{others.max():.3f}")
        if vi: issues[var] = vi
    return issues

def _apply_fixes_for_issues(series, issue_tags, original, seed, iteration):
    fixed = series.copy(); applied = []
    wp = max(0.01, 0.05 - (iteration-1)*0.01); js = 0.01*(1+(iteration-1)*0.5)
    has_zero = any("zero_variance" in t for t in issue_tags)
    has_out  = any("outliers"      in t for t in issue_tags)
    has_skew = any("skewness"      in t for t in issue_tags)
    has_kurt = any("kurtosis"      in t for t in issue_tags)
    has_coll = any("collinear"     in t for t in issue_tags)
    if has_zero: fixed = _add_jitter(fixed, seed=seed, scale=js); applied.append(f"Jitter (zero-var, iter {iteration})")
    if has_out:  fixed = _winsorize(fixed, limits=(wp, wp)); applied.append(f"Winsorize {wp*100:.0f}th–{(1-wp)*100:.0f}th (iter {iteration})")
    if has_skew:
        skw = float(fixed.skew())
        if skw > 2:   fixed = _log_transform(fixed);             applied.append(f"Log1p (pos skew {skw:.2f}, iter {iteration})")
        elif skw < -2: fixed = _log_transform(fixed.max()-fixed); applied.append(f"Reflected log (neg skew {skw:.2f}, iter {iteration})")
    if has_kurt and not has_out:
        kw = max(0.01, 0.025-(iteration-1)*0.005)
        fixed = _winsorize(fixed, limits=(kw,kw)); applied.append(f"Winsorize {kw*100:.1f}th (kurtosis, iter {iteration})")
    if has_coll: fixed = _add_jitter(fixed, seed=seed+10+iteration, scale=js); applied.append(f"Jitter (collinear, iter {iteration})")
    fixed = _rescale_to_original(fixed, original)
    return fixed, applied

def _efa_fix_pass(df_current, df_original, n_factors, rotation, load_thresh, comm_thresh, fix_history, seed, iteration):
    df_safe = _ensure_psd_dataframe(df_current)
    fa_res = run_efa(df_safe, n_factors, rotation)
    diag = diagnose_loadings(fa_res["loadings"], fa_res["communalities"], load_thresh, comm_thresh)
    still_flagged = diag[diag["RecommendDrop"]]["Variable"].tolist()
    if not still_flagged: return df_current, fa_res, diag, []
    df_next = df_current.copy(); data_issues = _detect_data_issues(df_current)
    for var in still_flagged:
        original_series = df_original[var]; current_series = df_current[var]
        raw_tags = data_issues.get(var, [])
        if not raw_tags:
            skw = float(current_series.skew()); krt = float(current_series.kurt())
            z = (current_series - current_series.mean()) / max(current_series.std(), 1e-9)
            n_out = int((z.abs() > 2.5).sum())
            if n_out > 0: raw_tags.append(f"outliers:{n_out}")
            if abs(skw) > 1.5: raw_tags.append(f"skewness:{skw:.3f}")
            if krt > 5: raw_tags.append(f"kurtosis:{krt:.3f}")
            if not raw_tags:
                others = [c for c in df_current.columns if c != var]
                residuals = current_series.copy()
                for ov in others:
                    cov = np.cov(residuals, df_current[ov])
                    if cov[1,1] > 1e-9: residuals = residuals - (cov[0,1]/cov[1,1])*df_current[ov]
                blend = 0.30*_rescale_to_original(residuals, current_series) + 0.70*current_series
                df_next[var] = _rescale_to_original(blend, original_series)
                fix_history.setdefault(var, []).append(f"Partial orthogonalisation blend (iter {iteration})")
                continue
        fixed_s, fix_desc = _apply_fixes_for_issues(current_series, raw_tags, original_series, seed=seed+iteration, iteration=iteration)
        df_next[var] = fixed_s; fix_history.setdefault(var, []).extend(fix_desc)
    for col in df_next.columns:
        if df_next[col].std() < 1e-9:
            df_next[col] = _add_jitter(df_original[col], seed=seed+hash(col)%1000, scale=0.05)
    return df_next, fa_res, diag, still_flagged

def run_auto_fix(df, initial_problem_vars, n_factors, rotation, load_thresh, comm_thresh, seed=42, max_iter=6):
    df_current = df.copy(); df_original = df.copy()
    fix_history = {}; iteration_log = []
    for iteration in range(1, max_iter+1):
        df_next, fa_res, diag, still_flagged = _efa_fix_pass(
            df_current, df_original, n_factors, rotation, load_thresh, comm_thresh, fix_history, seed, iteration)
        iteration_log.append(dict(Iteration=iteration, FlaggedVars=len(still_flagged),
                                   AvgCommunality=round(float(fa_res["communalities"].mean()),4),
                                   FixedThisPass=", ".join(still_flagged) if still_flagged else "—"))
        df_current = df_next
        if len(still_flagged) == 0: break
    df_final_safe = _ensure_psd_dataframe(df_current)
    final_fa = run_efa(df_final_safe, n_factors, rotation)
    final_diag = diagnose_loadings(final_fa["loadings"], final_fa["communalities"], load_thresh, comm_thresh)
    rows = []
    for var in df.columns.tolist():
        fixes = fix_history.get(var, [])
        di_str = " | ".join(_detect_data_issues(df).get(var, ["None"]))
        rows.append(dict(Variable=var, OriginalDataIssues=di_str,
                         FixesApplied=" → ".join(fixes) if fixes else "—",
                         Iterations=len(fixes),
                         FinalCommunality=round(float(final_fa["communalities"][var]), 4),
                         FinalIssue=final_diag.loc[final_diag["Variable"]==var,"Issue"].values[0]
                                    if var in final_diag["Variable"].values else "—",
                         Status="✔ Fixed" if var in fix_history else "✓ Clean"))
    return df_current, pd.DataFrame(rows), final_fa, final_diag, pd.DataFrame(iteration_log)

# ── CFA (preserved) ───────────────────────────────────────────────────────────
def build_cfa_model(loadings, threshold=0.4):
    factor_vars = {}
    for var in loadings.index:
        abs_row = np.abs(loadings.loc[var])
        if abs_row.max() >= threshold:
            factor_vars.setdefault(abs_row.idxmax(), []).append(var)
    factor_vars = {f: v for f, v in factor_vars.items() if len(v) >= 2}
    return "\n".join(f"{f} =~ " + " + ".join(v) for f, v in factor_vars.items()), factor_vars

def _parse_fit_indices(stats):
    flat = stats.iloc[0] if isinstance(stats, pd.DataFrame) else stats
    flat.index = [str(c).strip().upper() for c in flat.index]
    mapping = dict(CFI=["CFI"],TLI=["TLI","NNFI"],RMSEA=["RMSEA"],SRMR=["SRMR"],
                   Chi2=["CHI2","CHISQ","CHI-SQUARE","X2"],df=["DF","DOF"],
                   p_value=["P-VALUE","PVALUE","P_VALUE","P(CHI2)"],AIC=["AIC"],BIC=["BIC"])
    result = {}
    for label, candidates in mapping.items():
        for c in candidates:
            if c in flat.index:
                try: result[label] = round(float(flat[c]),4)
                except: result[label] = flat[c]
                break
    return result

def run_cfa(df, model_str):
    try:
        from semopy import Model; from semopy.stats import calc_stats
        model = Model(model_str); model.fit(df)
        try: fit_indices = _parse_fit_indices(calc_stats(model))
        except Exception as e: fit_indices = {"parse_error": str(e)}
        return dict(success=True, fit_indices=fit_indices, estimates=model.inspect(), model_str=model_str, error=None)
    except Exception as e:
        return dict(success=False, fit_indices={}, estimates=None, model_str=model_str, error=str(e))

def assess_cfa_fit(fit_indices, thresholds):
    assessment = {}
    for idx, direction in [("CFI","≥"),("TLI","≥"),("RMSEA","≤"),("SRMR","≤")]:
        if idx in fit_indices and idx in thresholds:
            val, thresh = fit_indices[idx], thresholds[idx]
            passed = val >= thresh if direction == "≥" else val <= thresh
            assessment[idx] = dict(value=val, threshold=thresh, pass_=passed, direction=direction)
    n_pass = sum(1 for v in assessment.values() if v["pass_"])
    return dict(indices=assessment, n_pass=n_pass, n_total=len(assessment),
                overall_pass=len(assessment)>0 and n_pass==len(assessment))

def get_modification_suggestions(fit_assessment):
    suggestions = []; indices = fit_assessment.get("indices", {})
    for idx, msg in [
        ("RMSEA","RMSEA exceeds threshold. Consider freeing residual covariances between items sharing method variance."),
        ("CFI",  "CFI is below threshold. Check whether indicators load on multiple factors."),
        ("SRMR", "SRMR exceeds threshold. Large residual correlations exist — review for systematic patterns."),
    ]:
        if idx in indices and not indices[idx]["pass_"]: suggestions.append(f"{idx} = {indices[idx]['value']:.3f} — {msg}")
    if not suggestions and not fit_assessment.get("overall_pass"):
        suggestions.append("Model fit is inadequate. Consider revising factor structure or removing low-communality items.")
    return suggestions

# ── Synthetic data (preserved) ────────────────────────────────────────────────
def _make_psd(matrix, eps=1e-8):
    ev, evec = np.linalg.eigh(matrix)
    return evec @ np.diag(np.maximum(ev, eps)) @ evec.T

def generate_factor_based(df, efa_result, n_samples=500, seed=42):
    np.random.seed(seed)
    loadings = efa_result["loadings"].values; communalities = efa_result["communalities"].values
    n_factors = efa_result["n_factors"]; columns = efa_result["loadings"].index.tolist()
    factor_scores = np.random.multivariate_normal(np.zeros(n_factors), np.eye(n_factors), n_samples)
    common = factor_scores @ loadings.T
    unique = np.random.normal(0, np.sqrt(np.maximum(1-communalities, 1e-6)), (n_samples, len(columns)))
    synthetic = common + unique
    orig_mean, orig_std = df[columns].mean().values, df[columns].std().values
    orig_std = np.where(orig_std==0, 1.0, orig_std)
    syn_mean, syn_std = synthetic.mean(0), synthetic.std(0); syn_std = np.where(syn_std==0, 1.0, syn_std)
    return pd.DataFrame(((synthetic - syn_mean) / syn_std) * orig_std + orig_mean, columns=columns)

def generate_correlation_based(df, n_samples=500, seed=42):
    np.random.seed(seed)
    cov = _make_psd(df.cov().values)
    return pd.DataFrame(np.random.multivariate_normal(df.mean().values, cov, n_samples), columns=df.columns.tolist())

def validate_synthetic(original, synthetic):
    return pd.DataFrame([dict(
        Variable=col, OrigMean=round(original[col].mean(),4), SynMean=round(synthetic[col].mean(),4),
        OrigStd=round(original[col].std(),4), SynStd=round(synthetic[col].std(),4),
        MeanDelta=round(abs(original[col].mean()-synthetic[col].mean()),4),
        StdDelta=round(abs(original[col].std()-synthetic[col].std()),4),
    ) for col in original.columns])

# ── Plot functions (preserved) ────────────────────────────────────────────────
def plot_scree(eigenvalues, suggested_n):
    x = list(range(1, len(eigenvalues)+1))
    colors = [C["green"] if i < suggested_n else C["muted"] for i in range(len(eigenvalues))]
    fig = go.Figure()
    fig.add_trace(go.Bar(x=x, y=eigenvalues, marker_color=colors, showlegend=False,
                         hovertemplate="Factor %{x}<br>λ = %{y:.3f}<extra></extra>"))
    fig.add_trace(go.Scatter(x=x, y=eigenvalues, mode="lines+markers",
                             line=dict(color=C["accent"], width=2),
                             marker=dict(size=7, color=C["accent"]), showlegend=False))
    fig.add_hline(y=1.0, line_dash="dash", line_color=C["yellow"],
                  annotation_text="Kaiser criterion (λ=1)", annotation_font_color=C["yellow"])
    fig.update_layout(**LAYOUT_BASE, height=360, showlegend=False,
                      title=dict(text=f"Scree Plot — Suggested factors: <b>{suggested_n}</b>",
                                 font=dict(color=C["accent2"])),
                      xaxis=dict(title="Factor", gridcolor=C["border"], tickmode="linear"),
                      yaxis=dict(title="Eigenvalue", gridcolor=C["border"]))
    return fig

def plot_loading_heatmap(loadings, threshold=0.4):
    z = loadings.values; variables, factors = loadings.index.tolist(), loadings.columns.tolist()
    annotations = [dict(x=f, y=v, text=f"{z[i][j]:.2f}", showarrow=False,
                        font=dict(color="white" if abs(z[i][j])>=threshold else C["muted"], size=11))
                   for i, v in enumerate(variables) for j, f in enumerate(factors)]
    fig = go.Figure(go.Heatmap(z=z, x=factors, y=variables, zmid=0, zmin=-1, zmax=1,
                               colorscale=[[0,"#0a0c24"],[0.5,C["surface"]],[1,C["accent"]]],
                               colorbar=dict(title="Loading", tickfont=dict(color=C["text"]),
                                             title_font=dict(color=C["text"])),
                               hovertemplate="%{y} → %{x}<br>Loading: %{z:.3f}<extra></extra>"))
    fig.update_layout(**LAYOUT_BASE, annotations=annotations,
                      title=dict(text="Factor Loading Heatmap", font=dict(color=C["accent2"])),
                      height=max(300, len(variables)*32+100), xaxis=dict(side="top"))
    return fig

def plot_communalities(communalities, comm_thresh=0.3):
    colors = [C["green"] if v>=0.5 else (C["yellow"] if v>=comm_thresh else C["red"])
              for v in communalities.values]
    fig = go.Figure(go.Bar(x=communalities.index.tolist(), y=communalities.values,
                           marker_color=colors,
                           hovertemplate="%{x}<br>Communality: %{y:.3f}<extra></extra>"))
    fig.add_hline(y=comm_thresh, line_dash="dash", line_color=C["red"],
                  annotation_text=f"Threshold ({comm_thresh})", annotation_font_color=C["red"])
    fig.add_hline(y=0.5, line_dash="dot", line_color=C["yellow"],
                  annotation_text="Good (0.50)", annotation_font_color=C["yellow"])
    fig.update_layout(**LAYOUT_BASE, height=360,
                      title=dict(text="Communalities per Variable", font=dict(color=C["accent2"])),
                      xaxis=dict(title="Variable", gridcolor=C["border"], tickangle=-35),
                      yaxis=dict(title="Communality", gridcolor=C["border"], range=[0,1]))
    return fig

def plot_fit_indices(fit_assessment):
    indices = fit_assessment.get("indices", {})
    if not indices: return go.Figure()
    labels = list(indices.keys()); values = [d["value"] for d in indices.values()]
    colors = [C["green"] if d["pass_"] else C["red"] for d in indices.values()]
    fig = go.Figure(go.Bar(x=labels, y=values, marker_color=colors,
                           text=[f"{v:.3f}" for v in values], textposition="outside",
                           hovertemplate="%{x}: %{y:.4f}<extra></extra>"))
    for label, data in indices.items():
        fig.add_annotation(x=label, y=data["threshold"],
                           text=f"Threshold: {data['threshold']}",
                           showarrow=True, arrowhead=2, arrowcolor=C["yellow"],
                           font=dict(color=C["yellow"], size=10), ay=-30)
    fig.update_layout(**LAYOUT_BASE, height=360, showlegend=False,
                      title=dict(text="CFA Fit Indices", font=dict(color=C["accent2"])),
                      xaxis=dict(title="Index", gridcolor=C["border"]),
                      yaxis=dict(title="Value", gridcolor=C["border"]))
    return fig

def plot_correlation_matrix(df):
    corr = df.corr(numeric_only=True).round(2); cols = corr.columns.tolist()
    annotations = [dict(x=c, y=r, text=f"{corr.loc[r,c]:.2f}", showarrow=False,
                        font=dict(size=9, color="white" if abs(corr.loc[r,c])>0.5 else C["text"]))
                   for r in corr.index for c in cols]
    fig = go.Figure(go.Heatmap(z=corr.values, x=cols, y=cols, colorscale="RdBu", zmid=0, zmin=-1, zmax=1,
                               colorbar=dict(title="r", tickfont=dict(color=C["text"]),
                                             title_font=dict(color=C["text"])),
                               hovertemplate="%{y} × %{x}<br>r = %{z:.2f}<extra></extra>"))
    fig.update_layout(**LAYOUT_BASE, annotations=annotations,
                      title=dict(text="Correlation Matrix", font=dict(color=C["accent2"])),
                      height=max(350, len(cols)*30+120), xaxis=dict(tickangle=-35))
    return fig

def plot_synthetic_comparison(original, synthetic, max_vars=6):
    cols = original.columns[:max_vars].tolist(); ncols = min(3, len(cols)); nrows = (len(cols)+ncols-1)//ncols
    fig = make_subplots(rows=nrows, cols=ncols, subplot_titles=cols)
    for i, col in enumerate(cols):
        r, ci = i//ncols+1, i%ncols+1
        fig.add_trace(go.Histogram(x=original[col], name="Original", nbinsx=20,
                                   marker_color=C["accent"], opacity=0.6, showlegend=(i==0)), row=r, col=ci)
        fig.add_trace(go.Histogram(x=synthetic[col], name="Synthetic", nbinsx=20,
                                   marker_color=C["accent2"], opacity=0.6, showlegend=(i==0)), row=r, col=ci)
    fig.update_layout(**LAYOUT_BASE, barmode="overlay", height=300*nrows,
                      title=dict(text="Original vs Synthetic Distributions", font=dict(color=C["accent2"])),
                      legend=dict(orientation="h", y=1.02, x=0.5, xanchor="center", font=dict(color=C["text"])))
    return fig

# ── DOCX Report (preserved) ───────────────────────────────────────────────────
def generate_docx_report(original_df, cleaned_df, suitability, efa_result,
                         diagnostics, dropped_vars, cfa_result, fit_assessment,
                         cfa_thresholds, synthetic_validation=None, model_str=""):
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor, Inches
    from docx.oxml.ns import qn; from docx.oxml import OxmlElement
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    doc = DocxDocument()
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Inches(1)
        section.left_margin = section.right_margin = Inches(1)

    def _rc(run, hx):
        run.font.color.rgb = RGBColor(int(hx[1:3],16), int(hx[3:5],16), int(hx[5:7],16))
    def _heading(text, level=1, color="#00e5c8"):
        p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text); run.bold = True; run.font.size = Pt(16 if level==1 else 13); _rc(run, color); return p
    def _para(text, bold=False, color=None, size=10):
        p = doc.add_paragraph(); run = p.add_run(text); run.bold = bold; run.font.size = Pt(size)
        if color: _rc(run, color); return p
    def _shade_cell(cell, hx):
        tc_pr = cell._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto"); shd.set(qn("w:fill"), hx.lstrip("#"))
        tc_pr.append(shd)
    def _table(headers, rows_data, header_bg="1f2535"):
        t = doc.add_table(rows=1, cols=len(headers)); t.style = "Table Grid"
        hdr_cells = t.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = h; run = hdr_cells[i].paragraphs[0].runs[0]
            run.bold = True; run.font.size = Pt(9); _rc(run, "#00e5c8"); _shade_cell(hdr_cells[i], header_bg)
        for row_vals, row_colors in rows_data:
            cells = t.add_row().cells
            for i, val in enumerate(row_vals):
                cells[i].text = str(val); run = cells[i].paragraphs[0].runs[0]; run.font.size = Pt(9)
                if row_colors and row_colors[i]: _rc(run, row_colors[i])
        return t

    ts = datetime.now().strftime("%Y-%m-%d %H:%M")
    title_p = doc.add_paragraph(); title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title_p.add_run("◈ EFActor — Analysis Report · Bayantx360 Suite")
    r.bold = True; r.font.size = Pt(20); _rc(r, "#00e5c8")
    meta_p = doc.add_paragraph(); meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = meta_p.add_run(f"Generated: {ts}"); mr.font.size = Pt(9); _rc(mr, "#6b7a9a")
    doc.add_paragraph()

    _heading("1. Dataset Overview")
    _table(["Metric","Value"],[
        (["Original Variables", str(len(original_df.columns))], [None,"#00e5c8"]),
        (["After Cleaning",     str(len(cleaned_df.columns))],  [None,"#00e5c8"]),
        (["Observations",       str(len(cleaned_df))],           [None,"#00e5c8"]),
        (["Dropped Variables",  str(len(dropped_vars))],         [None,"#f05c7c" if dropped_vars else "#22d3a0"]),
    ])
    if dropped_vars: _para(f"Dropped: {', '.join(dropped_vars)}", color="#f05c7c", size=9)

    _heading("2. EFA Suitability")
    s = suitability
    _table(["Test","Value","Threshold","Result"],[
        (["KMO Score", f"{s['kmo_model']} — {s['kmo_label']}", "≥ 0.60", "PASS" if s["kmo_pass"] else "FAIL"],
         [None,None,None,"#22d3a0" if s["kmo_pass"] else "#f05c7c"]),
        (["Bartlett's Sphericity", f"χ² = {s['bartlett_chi2']}, p = {s['bartlett_p']}", "p < 0.05", "PASS" if s["bartlett_pass"] else "FAIL"],
         [None,None,None,"#22d3a0" if s["bartlett_pass"] else "#f05c7c"]),
        (["Overall","","","PASS" if s["overall_pass"] else "FAIL"],
         [None,None,None,"#22d3a0" if s["overall_pass"] else "#f05c7c"]),
    ])

    _heading("3. Exploratory Factor Analysis")
    _para(f"Factors extracted: {efa_result['n_factors']}", bold=True, color="#00e5c8")
    loadings = efa_result["loadings"]; communalities = efa_result["communalities"]; variance = efa_result["variance"]
    _heading("Factor Loadings", level=2, color="#7c6df0")
    _table(["Variable"] + loadings.columns.tolist() + ["Communality"],
           [(([v] + [f"{loadings.loc[v,c]:.3f}" for c in loadings.columns] + [f"{communalities[v]:.3f}"]),
             ([None]*(len(loadings.columns)+1) + ["#22d3a0" if communalities[v]>=0.5 else ("#f5a623" if communalities[v]>=0.3 else "#f05c7c")]))
            for v in loadings.index])
    _heading("Variance Explained", level=2, color="#7c6df0")
    _table(["Factor","SS Loadings","Proportion Var","Cumulative Var"],
           [([idx,f"{r['SS Loadings']:.4f}",f"{r['Proportion Var']*100:.2f}%",f"{r['Cumulative Var']*100:.2f}%"],[None]*4)
            for idx, r in variance.iterrows()])
    _heading("Item Diagnostics", level=2, color="#7c6df0")
    diag_rows = []
    for _, r in diagnostics.iterrows():
        comm = r["Communality"]; issue = r["Issue"]
        cc = "#22d3a0" if comm>=0.5 else ("#f5a623" if comm>=0.3 else "#f05c7c")
        ic = "#22d3a0" if issue=="OK" else ("#f5a623" if "Cross" in issue else "#f05c7c")
        dc = "#f05c7c" if r["RecommendDrop"] else "#22d3a0"
        diag_rows.append(([r["Variable"],str(r["MaxLoading"]),str(r["FactorsAboveThreshold"]),str(comm),issue,"✗ Drop" if r["RecommendDrop"] else "✓ Keep"],[None,None,None,cc,ic,dc]))
    _table(["Variable","Max Load","# Factors ≥ Threshold","Communality","Issue","Recommended"], diag_rows)

    if cfa_result and cfa_result.get("success") and fit_assessment:
        _heading("4. Confirmatory Factor Analysis")
        fa = fit_assessment; overall = "ADEQUATE" if fa["overall_pass"] else "INADEQUATE"
        _para(f"Model Fit: {overall}  ({fa['n_pass']}/{fa['n_total']} indices passed)", bold=True,
              color="#22d3a0" if fa["overall_pass"] else "#f05c7c")
        _table(["Index","Value","Threshold","Result"],
               [([idx,str(d["value"]),f"{d['direction']} {d['threshold']}","PASS" if d["pass_"] else "FAIL"],
                 [None,None,None,"#22d3a0" if d["pass_"] else "#f05c7c"])
                for idx, d in fa["indices"].items()])
        _heading("Model Specification", level=2, color="#7c6df0"); doc.add_paragraph(model_str or "—")

    if synthetic_validation is not None:
        _heading("5. Synthetic Data Validation")
        _table(["Variable","Orig Mean","Syn Mean","Orig Std","Syn Std","Mean Δ","Std Δ"],
               [([r["Variable"],str(r["OrigMean"]),str(r["SynMean"]),str(r["OrigStd"]),str(r["SynStd"]),str(r["MeanDelta"]),str(r["StdDelta"])],[None]*7)
                for _, r in synthetic_validation.iterrows()])

    doc.add_paragraph(); foot_p = doc.add_paragraph(); foot_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = foot_p.add_run("EFActor · Bayantx360 Suite — Interpret results within your research design and theoretical framework.")
    fr.font.size = Pt(8); _rc(fr, "#6b7a9a")
    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf.getvalue()


# ═══════════════════════════════════════════════════════════════════════════════
# SIDEBAR
# ═══════════════════════════════════════════════════════════════════════════════

with st.sidebar:
    st.markdown("""
    <div style="padding:14px 0 20px 0;border-bottom:1px solid var(--border);margin-bottom:18px;">
        <div style="font-family:'Syne',sans-serif;font-size:1.3rem;font-weight:800;
                    color:var(--text);letter-spacing:-0.02em;">
            🔬 EF<span style="color:var(--accent2);">Act</span>or
        </div>
        <div style="font-family:'DM Mono',monospace;font-size:0.62rem;color:var(--muted);
                    margin-top:4px;letter-spacing:0.08em;">
            PSYCHOMETRIC ANALYSIS · BAYANTX360 SUITE
        </div>
    </div>
    """, unsafe_allow_html=True)

    render_credit_hud()

    if st.button("⬡ Back to Suite", use_container_width=True):
        st.switch_page(st.session_state["_home_page"])

    st.markdown("---")
    st.markdown("### ⚙️ EFA Settings")
    loading_threshold     = st.slider("Loading threshold",     0.30, 0.60, 0.40, 0.05)
    communality_threshold = st.slider("Communality threshold", 0.20, 0.60, 0.30, 0.05)
    rotation_method       = st.selectbox("Rotation method", ["varimax","oblimin","promax","quartimax","equamax"])

    st.markdown("---")
    st.markdown("### 📐 CFA Thresholds")
    cfi_thresh   = st.slider("CFI ≥",   0.80, 0.99, 0.95, 0.01)
    tli_thresh   = st.slider("TLI ≥",   0.80, 0.99, 0.95, 0.01)
    rmsea_thresh = st.slider("RMSEA ≤", 0.04, 0.15, 0.06, 0.01)
    srmr_thresh  = st.slider("SRMR ≤",  0.04, 0.15, 0.08, 0.01)
    cfa_thresholds = dict(CFI=cfi_thresh, TLI=tli_thresh, RMSEA=rmsea_thresh, SRMR=srmr_thresh)

    st.markdown("---")
    st.markdown("### 🧪 Synthetic Data")
    syn_n    = st.number_input("Sample size",  50, 10000, 500, 50)
    syn_seed = st.number_input("Random seed",   0,  9999,  42)

    st.markdown("---")
    if st.button("🔄 Reset Analysis", use_container_width=True):
        for k in ["_last_file_key","df_original","df_working","suitability","n_factors_auto",
                  "eigenvalues","efa_result","diagnostics","efa_done","dropped_vars",
                  "cfa_result","fit_assessment","synthetic_factor","synthetic_corr",
                  "syn_validation","report_docx","df_autofix","fix_log","autofix_efa_result"]:
            st.session_state.pop(k, None)
        st.rerun()

    if st.button("Sign Out", use_container_width=True):
        sign_out()

    if is_trial():
        st.markdown("---")
        st.markdown("""
        <div style="background:rgba(0,229,200,0.05);border:1px solid rgba(0,229,200,0.2);
                    border-radius:10px;padding:14px 16px;text-align:center;">
            <div style="font-family:'Syne',sans-serif;font-size:0.82rem;font-weight:700;
                        color:var(--accent);margin-bottom:6px;">Unlock Exports</div>
            <div style="font-family:'DM Mono',monospace;font-size:0.62rem;color:var(--muted);
                        line-height:1.7;margin-bottom:10px;">
                Get credits to download datasets and reports.
            </div>
        </div>
        """, unsafe_allow_html=True)
        st.link_button("Get Access Key →", "https://x.com/bayantx360", use_container_width=True)


# ═══════════════════════════════════════════════════════════════════════════════
# MAIN CONTENT
# ═══════════════════════════════════════════════════════════════════════════════

st.markdown("""
<div class="app-hero">
    <div class="app-hero-title">EF<span>Act</span>or</div>
    <div class="app-hero-sub">🔬 Psychometric Analysis Platform · Bayantx360 Suite</div>
</div>
""", unsafe_allow_html=True)

st.markdown("_Upload your dataset and run a complete EFA → CFA pipeline. Diagnose items, confirm structure, generate synthetic data, and export your results._")

if is_trial():
    st.markdown(f"""
    <div style="background:rgba(245,166,35,0.07);border:1px solid rgba(245,166,35,0.2);
                border-left:3px solid var(--warn);border-radius:8px;padding:12px 16px;
                font-family:'DM Mono',monospace;font-size:0.72rem;color:var(--warn);line-height:1.7;margin-bottom:16px;">
        ⚡ <strong style="color:var(--text);">Free Trial mode</strong> — all analysis runs are unlimited.
        Data export &amp; DOCX report require a paid access key.
        <a href="https://x.com/bayantx360" style="color:var(--accent);margin-left:4px;">Get access →</a>
    </div>
    """, unsafe_allow_html=True)

st.markdown("---")

# ── STEP 1: UPLOAD ─────────────────────────────────────────────────────────────
st.markdown('<div class="scard-title">Step 1 · Upload Dataset</div>', unsafe_allow_html=True)
col_up1, col_up2 = st.columns([2, 1])
with col_up1:
    uploaded = st.file_uploader("Upload CSV or Excel", type=["csv","CSV","xlsx","xls","XLSX","XLS"],
                                 help="Numeric columns only. Missing values auto-dropped.")
with col_up2:
    st.markdown("""
    <div class="scard" style="font-size:0.7rem;line-height:1.9;color:var(--muted);">
        <strong style="color:var(--text);">Requirements</strong><br>
        • Numeric variables only<br>
        • Minimum 5 variables<br>
        • Recommended ≥ 100 rows<br>
        • Missing values auto-dropped
    </div>
    """, unsafe_allow_html=True)

if uploaded:
    _file_key = f"{uploaded.name}_{uploaded.size}"
    if S.get("_last_file_key") != _file_key:
        try:
            fname = uploaded.name.lower()
            if fname.endswith(".csv"):
                try: df_raw = pd.read_csv(uploaded)
                except UnicodeDecodeError: uploaded.seek(0); df_raw = pd.read_csv(uploaded, encoding="latin-1")
            elif fname.endswith((".xlsx",".xls")):
                df_raw = pd.read_excel(uploaded)
            else:
                st.error("❌ Unsupported file type."); st.stop()
            df_coerced = df_raw.apply(lambda col: pd.to_numeric(col, errors="coerce"))
            df_numeric = df_coerced.select_dtypes(include=[np.number]).dropna(axis=1, how="all").dropna()
            if len(df_numeric.columns) < 3:
                st.error("❌ Need at least 3 numeric variables."); st.stop()
            if len(df_numeric) < 30:
                st.warning("⚠️ Fewer than 30 observations — results may be unreliable.")
            S.df_original = df_numeric.copy(); S.df_working = df_numeric.copy(); S.dropped_vars = []
            for k in ["suitability","efa_result","cfa_result","fit_assessment","efa_done",
                      "synthetic_factor","synthetic_corr","syn_validation","report_docx",
                      "n_factors_auto","eigenvalues","diagnostics"]:
                S[k] = None if k != "efa_done" else False
            S["_last_file_key"] = _file_key
        except Exception as e:
            st.error(f"❌ Could not parse file: {e}"); st.stop()

if S.df_original is None:
    st.markdown("""
    <div style="background:rgba(245,166,35,0.06);border:1px solid rgba(245,166,35,0.2);
                border-radius:8px;padding:12px 16px;font-family:'DM Mono',monospace;
                font-size:0.72rem;color:var(--warn);margin:12px 0;">
        👆 Upload a dataset to begin.
    </div>
    """, unsafe_allow_html=True)
    st.stop()

with st.expander("📋 Dataset Preview", expanded=True):
    c1, c2, c3, c4 = st.columns(4)
    for col, val, label in [
        (c1, len(S.df_original), "Observations"), (c2, len(S.df_original.columns), "Variables"),
        (c3, len(S.df_working.columns), "Working Vars"), (c4, len(S.dropped_vars), "Dropped"),
    ]:
        col.markdown(f"""
        <div class="scard" style="text-align:center;padding:16px 12px;">
            <div style="font-family:'Syne',sans-serif;font-size:1.6rem;font-weight:800;color:var(--accent);">{val}</div>
            <div style="font-family:'DM Mono',monospace;font-size:0.6rem;text-transform:uppercase;letter-spacing:0.12em;color:var(--muted);margin-top:4px;">{label}</div>
        </div>
        """, unsafe_allow_html=True)
    st.dataframe(S.df_working.head(10), use_container_width=True)
    t1, t2 = st.tabs(["Descriptive Statistics", "Correlation Matrix"])
    with t1: st.dataframe(S.df_working.describe().T.round(4), use_container_width=True)
    with t2: st.plotly_chart(plot_correlation_matrix(S.df_working), use_container_width=True)

st.markdown("---")

# ── STEP 2: SUITABILITY ────────────────────────────────────────────────────────
st.markdown('<div class="scard-title">Step 2 · EFA Suitability Tests</div>', unsafe_allow_html=True)
if st.button("▶ Run Suitability Tests"):
    with st.spinner("Running KMO and Bartlett's tests…"):
        S.suitability = check_efa_suitability(S.df_working)
        ev_info = determine_n_factors(S.df_working)
        S.n_factors_auto = ev_info["suggested_n"]; S.eigenvalues = ev_info["eigenvalues"]

if S.suitability:
    suit = S.suitability
    overall_style = "color:var(--success)" if suit["overall_pass"] else "color:var(--accent3)"
    overall_text  = "✓ SUITABLE FOR EFA" if suit["overall_pass"] else "✗ EFA SUITABILITY ISSUES"
    st.markdown(f'<div style="font-family:\'DM Mono\',monospace;font-size:0.8rem;{overall_style};margin-bottom:10px;font-weight:700;">{overall_text}</div>', unsafe_allow_html=True)
    c1, c2 = st.columns(2)
    with c1:
        color_kmo = "var(--success)" if suit["kmo_pass"] else "var(--accent3)"
        st.markdown(f"""
        <div class="scard">
            <div class="scard-title">KMO Score</div>
            <div style="font-family:'Syne',sans-serif;font-size:1.6rem;font-weight:800;color:{color_kmo};">{suit['kmo_model']}</div>
            <div style="font-family:'DM Mono',monospace;font-size:0.62rem;color:var(--muted);margin-top:4px;">{suit['kmo_label']} · {'PASS' if suit['kmo_pass'] else 'FAIL'}</div>
        </div>
        """, unsafe_allow_html=True)
    with c2:
        color_bt = "var(--success)" if suit["bartlett_pass"] else "var(--accent3)"
        st.markdown(f"""
        <div class="scard">
            <div class="scard-title">Bartlett p-value (χ²={suit['bartlett_chi2']})</div>
            <div style="font-family:'Syne',sans-serif;font-size:1.6rem;font-weight:800;color:{color_bt};">{suit['bartlett_p']}</div>
            <div style="font-family:'DM Mono',monospace;font-size:0.62rem;color:var(--muted);margin-top:4px;">{'PASS' if suit['bartlett_pass'] else 'FAIL'}</div>
        </div>
        """, unsafe_allow_html=True)
    if not suit["overall_pass"]:
        st.warning("⚠️ Suitability concerns detected. You can still proceed — interpret results cautiously.")
    st.markdown("#### Scree Plot")
    st.plotly_chart(plot_scree(S.eigenvalues, S.n_factors_auto), use_container_width=True)
    st.info(f"🔢 Kaiser criterion suggests **{S.n_factors_auto}** factor(s). Override below if theory or scree plot suggests otherwise.")

st.markdown("---")

# ── STEP 3: EFA ────────────────────────────────────────────────────────────────
st.markdown('<div class="scard-title">Step 3 · Exploratory Factor Analysis (EFA)</div>', unsafe_allow_html=True)
if S.suitability is None:
    st.warning("Complete Step 2 first.")
else:
    _max_factors = max(2, len(S.df_working.columns)-1)
    default_n = int(min(max(1, S.n_factors_auto if S.n_factors_auto else 2), _max_factors))
    n_factors = st.number_input("Number of factors to extract", 1, _max_factors, default_n, 1)
    if st.button("▶ Run EFA", use_container_width=True):
        with st.spinner("Running factor analysis…"):
            S.efa_result  = run_efa(S.df_working, n_factors, rotation_method)
            S.diagnostics = diagnose_loadings(S.efa_result["loadings"], S.efa_result["communalities"],
                                              loading_threshold, communality_threshold)
            S.efa_done = True
    if S.efa_result:
        t1, t2, t3, t4 = st.tabs(["Factor Loadings","Variance Explained","Loading Heatmap","Communalities"])
        with t1: st.dataframe(S.efa_result["loadings"].round(4), use_container_width=True)
        with t2: st.dataframe(S.efa_result["variance"].round(4), use_container_width=True)
        with t3: st.plotly_chart(plot_loading_heatmap(S.efa_result["loadings"], loading_threshold), use_container_width=True)
        with t4: st.plotly_chart(plot_communalities(S.efa_result["communalities"], communality_threshold), use_container_width=True)

st.markdown("---")

# ── STEP 4: DIAGNOSE & AUTO-FIX ───────────────────────────────────────────────
st.markdown('<div class="scard-title">Step 4 · Diagnose, Auto-Fix &amp; Refine Items</div>', unsafe_allow_html=True)
if S.efa_result is None:
    st.warning("Complete EFA (Step 3) first.")
else:
    # Colour-coded diagnostics table (preserved HTML rendering)
    def _diag_table_html(df):
        rows_html = ""
        for _, r in df.iterrows():
            comm = r["Communality"]
            cc = ("color:#22d3a0;font-weight:600;" if comm >= 0.5 else
                  ("color:#f5a623;font-weight:600;" if comm >= 0.3 else "color:#f05c7c;font-weight:600;"))
            ci = "🟢" if comm >= 0.5 else ("🟡" if comm >= 0.3 else "🔴")
            issue = r["Issue"]
            ic = ("color:#22d3a0;font-weight:600;" if issue=="OK" else
                  ("color:#f5a623;font-weight:600;" if "Cross" in issue else "color:#f05c7c;font-weight:600;"))
            ii = f"✓ {issue}" if issue=="OK" else (f"⚠ {issue}" if "Cross" in issue else f"✗ {issue}")
            drop = r["RecommendDrop"]
            dc = "color:#f05c7c;font-weight:700;" if drop else "color:#22d3a0;font-weight:600;"
            dd = "✗ Drop" if drop else "✓ Keep"
            rows_html += (f"<tr>"
                          f"<td style='padding:7px 10px;border-bottom:1px solid var(--border);'>{r['Variable']}</td>"
                          f"<td style='padding:7px 10px;border-bottom:1px solid var(--border);'>{r['MaxLoading']}</td>"
                          f"<td style='padding:7px 10px;border-bottom:1px solid var(--border);text-align:center;'>{r['FactorsAboveThreshold']}</td>"
                          f"<td style='padding:7px 10px;border-bottom:1px solid var(--border);{cc}'>{ci} {comm}</td>"
                          f"<td style='padding:7px 10px;border-bottom:1px solid var(--border);{ic}'>{ii}</td>"
                          f"<td style='padding:7px 10px;border-bottom:1px solid var(--border);{dc}'>{dd}</td>"
                          f"</tr>")
        HDR = "padding:8px 10px;text-align:left;color:var(--accent);font-family:'DM Mono',monospace;font-size:0.62rem;text-transform:uppercase;letter-spacing:0.1em;"
        return f"""<div style="overflow-x:auto;"><table style="width:100%;border-collapse:collapse;font-size:.85rem;">
<thead><tr style="background:var(--surface2);">
<th style="{HDR}">Variable</th><th style="{HDR}">Max Loading</th>
<th style="{HDR};text-align:center;"># Factors ≥ Threshold</th>
<th style="{HDR}">Communality</th><th style="{HDR}">Issue</th><th style="{HDR}">Recommended</th>
</tr></thead><tbody style="color:var(--text);">{rows_html}</tbody></table></div>"""

    st.markdown(_diag_table_html(S.diagnostics), unsafe_allow_html=True)
    problem_vars = S.diagnostics[S.diagnostics["RecommendDrop"]]["Variable"].tolist()
    if problem_vars:
        st.warning(f"⚠️ **{len(problem_vars)}** variable(s) flagged: {', '.join(problem_vars)}")

    st.markdown("---")
    st.markdown("""
    <div style="background:rgba(34,211,160,0.05);border:1px solid rgba(34,211,160,0.22);
                border-left:3px solid var(--success);border-radius:8px;padding:16px 20px;margin-bottom:18px;">
        <div style="font-family:'Syne',sans-serif;font-size:0.95rem;font-weight:700;color:var(--success);margin-bottom:6px;">
            ⚡ EFA Auto-Fix Engine
        </div>
        <div style="font-family:'DM Mono',monospace;font-size:0.7rem;color:var(--muted);line-height:1.8;">
            Automatically detects data-level causes of EFA problems (outliers, skewness, kurtosis,
            near-zero variance, collinearity) and applies targeted fixes —
            <strong style="color:var(--text);">without dropping any variable</strong>.
        </div>
    </div>
    """, unsafe_allow_html=True)

    col_fix1, col_fix2 = st.columns([2, 1])
    with col_fix1:
        if st.button("⚡ Run Auto-Fix (keep all variables)", use_container_width=True, key="autofix_btn"):
            with st.spinner("Running iterative EFA-aware fix engine (up to 6 passes)…"):
                df_fixed, fix_log_df, final_fa, final_diag, iter_log_df = run_auto_fix(
                    S.df_working, initial_problem_vars=problem_vars,
                    n_factors=S.efa_result["n_factors"], rotation=rotation_method,
                    load_thresh=loading_threshold, comm_thresh=communality_threshold, seed=42, max_iter=6)
                S.df_autofix = df_fixed; S.fix_log = fix_log_df
                S.autofix_efa_result = dict(efa=final_fa, diag=final_diag, iter_log=iter_log_df)
            n_remaining = int(final_diag["RecommendDrop"].sum())
            if n_remaining == 0:
                st.success(f"✓ All variables fixed. All {len(df_fixed.columns)} variables pass EFA diagnostics.")
            else:
                st.warning(f"⚠️ {n_remaining} variable(s) still flagged after 6 iterations — structural issues. Best-possible dataset still returned.")

    with col_fix2:
        if S.df_autofix is not None:
            if st.button("✅ Use Fixed Dataset for CFA / Export", use_container_width=True, key="apply_fix_btn"):
                S.df_working = S.df_autofix.copy()
                with st.spinner("Updating EFA on fixed dataset…"):
                    S.efa_result  = run_efa(S.df_working, S.efa_result["n_factors"], rotation_method)
                    S.diagnostics = diagnose_loadings(S.efa_result["loadings"], S.efa_result["communalities"],
                                                      loading_threshold, communality_threshold)
                    S.efa_done = True
                st.success("✓ Working dataset updated with fixes. Proceed to Step 5."); st.rerun()

    if S.fix_log is not None:
        fix_df = S.fix_log
        n_fixed = int((fix_df["Status"]=="✔ Fixed").sum())
        n_clean = int((fix_df["Status"]=="✓ Clean").sum())
        mc1, mc2, mc3 = st.columns(3)
        for col, val, label, color in [
            (mc1, n_fixed, "Variables Fixed", "var(--success)"),
            (mc2, n_clean, "Already Clean",   "var(--accent2)"),
            (mc3, len(fix_df), "Total Retained", "var(--text)"),
        ]:
            col.markdown(f"""
            <div class="scard" style="text-align:center;padding:16px;">
                <div style="font-family:'Syne',sans-serif;font-size:1.6rem;font-weight:800;color:{color};">{val}</div>
                <div style="font-family:'DM Mono',monospace;font-size:0.6rem;text-transform:uppercase;letter-spacing:0.1em;color:var(--muted);margin-top:4px;">{label}</div>
            </div>
            """, unsafe_allow_html=True)

        if S.autofix_efa_result and "iter_log" in S.autofix_efa_result:
            iter_df = S.autofix_efa_result["iter_log"]
            if len(iter_df) > 0:
                st.markdown("#### 🔁 Iterative Fix Progress")
                iter_rows = ""
                for _, r in iter_df.iterrows():
                    flagged = int(r["FlaggedVars"])
                    fc = "color:#22d3a0;" if flagged==0 else ("color:#f5a623;" if flagged<=2 else "color:#f05c7c;")
                    iter_rows += (f"<tr>"
                                  f"<td style='padding:6px 10px;border-bottom:1px solid var(--border);color:var(--text);text-align:center;'>Pass {int(r['Iteration'])}</td>"
                                  f"<td style='padding:6px 10px;border-bottom:1px solid var(--border);{fc}font-weight:700;text-align:center;'>{flagged}</td>"
                                  f"<td style='padding:6px 10px;border-bottom:1px solid var(--border);color:var(--accent2);text-align:center;'>{r['AvgCommunality']}</td>"
                                  f"<td style='padding:6px 10px;border-bottom:1px solid var(--border);color:var(--accent);font-size:.82rem;'>{r['FixedThisPass']}</td>"
                                  f"</tr>")
                HDR = "padding:8px 10px;color:var(--accent);font-family:'DM Mono',monospace;font-size:0.62rem;text-transform:uppercase;letter-spacing:0.1em;text-align:center;"
                st.markdown(f"""<div style="overflow-x:auto;"><table style="width:100%;border-collapse:collapse;font-size:.83rem;">
<thead><tr style="background:var(--surface2);">
<th style="{HDR}">Pass</th><th style="{HDR}">Still Flagged</th>
<th style="{HDR}">Avg Communality</th><th style="{HDR};text-align:left;">Variables Worked On</th>
</tr></thead><tbody>{iter_rows}</tbody></table></div>""", unsafe_allow_html=True)

        st.markdown("#### 🔧 Fix Log — Variable-by-Variable")
        fix_rows = ""
        for _, r in fix_df.iterrows():
            status = r["Status"]
            sc = "color:#22d3a0;font-weight:700;" if "Fixed" in status else "color:var(--muted);"
            si = "✔" if "Fixed" in status else "✓"
            fi = str(r.get("FinalIssue","—"))
            ic = "color:#22d3a0;" if fi in ("OK","—") else "color:#f5a623;"
            fix_rows += (f"<tr>"
                         f"<td style='padding:7px 10px;border-bottom:1px solid var(--border);color:var(--text);font-weight:600;'>{r['Variable']}</td>"
                         f"<td style='padding:7px 10px;border-bottom:1px solid var(--border);color:var(--accent3);font-size:.82rem;'>{r['OriginalDataIssues']}</td>"
                         f"<td style='padding:7px 10px;border-bottom:1px solid var(--border);color:var(--accent2);font-size:.82rem;'>{r['FixesApplied']}</td>"
                         f"<td style='padding:7px 10px;border-bottom:1px solid var(--border);{ic}font-size:.82rem;'>{fi}</td>"
                         f"<td style='padding:7px 10px;border-bottom:1px solid var(--border);color:var(--accent);text-align:center;'>{r['FinalCommunality']}</td>"
                         f"<td style='padding:7px 10px;border-bottom:1px solid var(--border);{sc}'>{si} {status.replace('✔ ','').replace('✓ ','')}</td>"
                         f"</tr>")
        HDR2 = "padding:8px 10px;text-align:left;color:var(--accent);font-family:'DM Mono',monospace;font-size:0.62rem;text-transform:uppercase;letter-spacing:0.1em;"
        st.markdown(f"""<div style="overflow-x:auto;"><table style="width:100%;border-collapse:collapse;font-size:.83rem;">
<thead><tr style="background:var(--surface2);">
<th style="{HDR2}">Variable</th><th style="{HDR2}">Raw Data Issues</th>
<th style="{HDR2}">Fixes Applied</th><th style="{HDR2}">Final EFA Issue</th>
<th style="{HDR2};text-align:center;">Final Communality</th><th style="{HDR2}">Status</th>
</tr></thead><tbody>{fix_rows}</tbody></table></div>""", unsafe_allow_html=True)

        if S.autofix_efa_result:
            st.markdown("#### 📊 EFA Comparison: Before vs After Fix")
            efa_b = S.efa_result; efa_a = S.autofix_efa_result["efa"]
            diag_b = S.diagnostics; diag_a = S.autofix_efa_result["diag"]
            nb = int(diag_b["RecommendDrop"].sum()); na = int(diag_a["RecommendDrop"].sum())
            cb = round(float(efa_b["communalities"].mean()),3); ca = round(float(efa_a["communalities"].mean()),3)
            delta = round(ca-cb,3); delta_str = f"+{delta}" if delta>=0 else str(delta)
            delta_color = "var(--success)" if delta>=0 else "var(--accent3)"
            cc1, cc2, cc3, cc4 = st.columns(4)
            for col, val, label, color in [
                (cc1,nb,"Flagged (Before)","var(--accent3)"),
                (cc2,na,"Flagged (After)","var(--success)"),
                (cc3,ca,"Avg Communality (After)","var(--accent2)"),
                (cc4,delta_str,"Communality Δ",delta_color),
            ]:
                col.markdown(f"""
                <div class="scard" style="text-align:center;padding:14px;">
                    <div style="font-family:'Syne',sans-serif;font-size:1.4rem;font-weight:800;color:{color};">{val}</div>
                    <div style="font-family:'DM Mono',monospace;font-size:0.58rem;text-transform:uppercase;letter-spacing:0.1em;color:var(--muted);margin-top:4px;">{label}</div>
                </div>
                """, unsafe_allow_html=True)
            t_before, t_after = st.tabs(["Communalities: Before", "Communalities: After"])
            with t_before: st.plotly_chart(plot_communalities(efa_b["communalities"], communality_threshold), use_container_width=True, key="comm_b")
            with t_after:  st.plotly_chart(plot_communalities(efa_a["communalities"], communality_threshold), use_container_width=True, key="comm_a")

        if S.df_autofix is not None:
            st.markdown("#### 📋 Fixed Dataset Preview")
            st.dataframe(S.df_autofix.head(10).round(4), use_container_width=True)
            st.info(f"✅ **{len(S.df_autofix.columns)} variables × {len(S.df_autofix)} rows** — all variables retained, problematic ones transformed in-place.")
            st.download_button("⬇ Download Fixed Dataset (CSV)",
                               data=S.df_autofix.to_csv(index=False).encode("utf-8"),
                               file_name="efactor_autofix_dataset.csv", mime="text/csv",
                               use_container_width=True, key="dl_fixed_inline")

    st.markdown("---")
    st.markdown("**— or — manually drop variables (traditional approach)**")
    available = S.df_working.columns.tolist()
    to_drop = st.multiselect("Select variables to drop (optional)", available,
                              default=[v for v in problem_vars if v in available])
    col_d1, col_d2 = st.columns(2)
    with col_d1:
        if st.button("▶ Apply Drops & Re-run EFA", use_container_width=True):
            if to_drop:
                S.df_working = S.df_working.drop(columns=to_drop)
                S.dropped_vars = list(set(S.dropped_vars + to_drop))
            with st.spinner("Re-running EFA…"):
                S.efa_result = run_efa(S.df_working, n_factors, rotation_method)
                S.diagnostics = diagnose_loadings(S.efa_result["loadings"], S.efa_result["communalities"],
                                                  loading_threshold, communality_threshold)
            st.success(f"✓ EFA re-run. Working set: {len(S.df_working.columns)} variables.")
    with col_d2:
        if st.button("↩ Restore All Variables", use_container_width=True):
            S.df_working = S.df_original.copy(); S.dropped_vars = []; S.efa_result = None
            S.diagnostics = None; S.efa_done = False; S.df_autofix = None
            S.fix_log = None; S.autofix_efa_result = None; st.rerun()
    if S.dropped_vars:
        st.info(f"📋 Dropped so far: **{', '.join(S.dropped_vars)}**")

st.markdown("---")

# ── STEP 5: CFA ────────────────────────────────────────────────────────────────
st.markdown('<div class="scard-title">Step 5 · Confirmatory Factor Analysis (CFA)</div>', unsafe_allow_html=True)
if S.efa_result is None:
    st.warning("Complete EFA (Step 3) first.")
else:
    model_str, factor_vars = build_cfa_model(S.efa_result["loadings"], loading_threshold)
    with st.expander("📝 Auto-generated CFA Model", expanded=False):
        edited_model = st.text_area("Edit model specification if needed:", value=model_str, height=160)
    if st.button("▶ Run CFA", use_container_width=True):
        if not edited_model.strip():
            st.error("Model specification is empty.")
        else:
            with st.spinner("Fitting CFA model…"):
                S.cfa_result = run_cfa(S.df_working, edited_model)
                if S.cfa_result["success"]:
                    S.fit_assessment = assess_cfa_fit(S.cfa_result["fit_indices"], cfa_thresholds)
                    S.cfa_result["model_str"] = edited_model
    if S.cfa_result:
        if not S.cfa_result["success"]:
            st.error(f"CFA failed: {S.cfa_result['error']}")
        else:
            fa = S.fit_assessment
            fit_style = "color:var(--success)" if fa["overall_pass"] else "color:var(--accent3)"
            fit_text  = "✓ MODEL FIT ADEQUATE" if fa["overall_pass"] else "✗ MODEL FIT INADEQUATE"
            st.markdown(f'<div style="font-family:\'DM Mono\',monospace;font-size:0.8rem;{fit_style};margin-bottom:10px;font-weight:700;">{fit_text} — {fa["n_pass"]}/{fa["n_total"]} indices passed</div>', unsafe_allow_html=True)
            st.plotly_chart(plot_fit_indices(fa), use_container_width=True)
            fit_records = [dict(Index=idx, Value=d["value"], Threshold=f"{d['direction']} {d['threshold']}",
                                Status="✓ PASS" if d["pass_"] else "✗ FAIL")
                           for idx, d in fa["indices"].items()]
            st.dataframe(pd.DataFrame(fit_records), use_container_width=True, hide_index=True)
            with st.expander("📊 Parameter Estimates"):
                if S.cfa_result["estimates"] is not None:
                    st.dataframe(S.cfa_result["estimates"], use_container_width=True)
            if not fa["overall_pass"]:
                st.markdown("#### 🔧 Modification Suggestions")
                for sg in get_modification_suggestions(fa):
                    st.warning(f"⚠️ {sg}")
                st.info("💡 Return to Step 4 to drop additional items, or Step 3 to re-extract with different n_factors, then re-run CFA.")

st.markdown("---")

# ── STEP 6: SYNTHETIC DATA ─────────────────────────────────────────────────────
st.markdown('<div class="scard-title">Step 6 · Synthetic Data Generation</div>', unsafe_allow_html=True)
if S.efa_result is None:
    st.warning("Complete EFA (Step 3) first.")
else:
    c1, c2 = st.columns(2)
    with c1:
        st.markdown("#### Factor-Structure Preserving")
        st.info("Simulates latent factor scores × loadings + unique variance. Preserves **psychometric structure**. Recommended for structural validity studies.")
        if st.button("▶ Generate (Factor-Based)", use_container_width=True):
            with st.spinner("Simulating from factor structure…"):
                S.synthetic_factor = generate_factor_based(S.df_working, S.efa_result, n_samples=syn_n, seed=syn_seed)
                S.syn_validation = validate_synthetic(S.df_working, S.synthetic_factor)
            st.success(f"✓ {syn_n} synthetic observations generated.")
    with c2:
        st.markdown("#### Correlation Preserving")
        st.info("Multivariate normal from empirical **covariance matrix**. Faster, preserves pairwise correlations but not latent structure explicitly.")
        if st.button("▶ Generate (Correlation-Based)", use_container_width=True):
            with st.spinner("Sampling from multivariate normal…"):
                S.synthetic_corr = generate_correlation_based(S.df_working, n_samples=syn_n, seed=syn_seed)
                S.syn_validation = validate_synthetic(S.df_working, S.synthetic_corr)
            st.success(f"✓ {syn_n} synthetic observations generated.")

    syn_display = S.synthetic_factor if S.synthetic_factor is not None else S.synthetic_corr
    if syn_display is not None:
        t1, t2, t3 = st.tabs(["Preview","Validation Summary","Distribution Comparison"])
        with t1: st.dataframe(syn_display.head(10).round(3), use_container_width=True)
        with t2:
            if S.syn_validation is not None:
                st.dataframe(S.syn_validation, use_container_width=True, hide_index=True)
        with t3: st.plotly_chart(plot_synthetic_comparison(S.df_working, syn_display), use_container_width=True)

st.markdown("---")

# ── STEP 7: EXPORT ─────────────────────────────────────────────────────────────
st.markdown('<div class="scard-title">Step 7 · Export Results</div>', unsafe_allow_html=True)

if S.efa_result is None:
    st.warning("Complete EFA (Step 3) to enable exports.")
else:
    syn_export = S.synthetic_factor if S.synthetic_factor is not None else S.synthetic_corr
    syn_label  = "factor_based" if S.synthetic_factor is not None else "correlation_based"
    REPORT_COST = 1
    current_credits = st.session_state.user_credits

    if is_trial():
        render_locked_banner("Data Export & DOCX Report", is_trial_user=True)
        _, uc, _ = st.columns([1, 2, 1])
        with uc:
            st.link_button("◈ Get Access Key →", "https://x.com/bayantx360", use_container_width=True)
    else:
        col_e1, col_e2, col_e3 = st.columns(3)

        with col_e1:
            st.markdown("##### 🗃️ Cleaned Dataset")
            n_clean = len(S.df_working); cost_clean = export_credit_cost(n_clean)
            st.info(f"{n_clean} rows × {len(S.df_working.columns)} cols  **Cost: {cost_clean} credit{'s' if cost_clean>1 else ''}**")
            if current_credits < cost_clean:
                st.warning(f"⚠️ Need {cost_clean} credits ({current_credits} remaining).")
                st.link_button("Top up →", "https://x.com/bayantx360", use_container_width=True)
            else:
                if st.button(f"⬇ Download Cleaned Data ({cost_clean} cr)", use_container_width=True, key="dl_clean_btn"):
                    handle_credit_deduction(cost_clean, app="EFActor", action="Cleaned Data Export")
                    st.download_button("⬇ cleaned_data.csv",
                                       data=S.df_working.to_csv(index=False).encode("utf-8"),
                                       file_name="efactor_cleaned_data.csv", mime="text/csv",
                                       use_container_width=True, key="dl_clean_actual")
                    st.success(f"✓ {cost_clean} credit(s) deducted.")

        with col_e2:
            st.markdown("##### 🧪 Synthetic Dataset")
            if syn_export is not None:
                n_syn = len(syn_export); cost_syn = export_credit_cost(n_syn)
                st.info(f"{n_syn} rows × {len(syn_export.columns)} cols  **Cost: {cost_syn} credit{'s' if cost_syn>1 else ''}**")
                if current_credits < cost_syn:
                    st.warning(f"⚠️ Need {cost_syn} credits ({current_credits} remaining).")
                    st.link_button("Top up →", "https://x.com/bayantx360", use_container_width=True)
                else:
                    if st.button(f"⬇ Download Synthetic Data ({cost_syn} cr)", use_container_width=True, key="dl_syn_btn"):
                        handle_credit_deduction(cost_syn, app="EFActor", action="Synthetic Data Export")
                        st.download_button(f"⬇ synthetic_{syn_label}.csv",
                                           data=syn_export.to_csv(index=False).encode("utf-8"),
                                           file_name=f"efactor_synthetic_{syn_label}.csv", mime="text/csv",
                                           use_container_width=True, key="dl_syn_actual")
                        st.success(f"✓ {cost_syn} credit(s) deducted.")
            else:
                st.warning("Generate synthetic data in Step 6 first.")

        with col_e3:
            st.markdown("##### 📄 Analysis Report")
            has_cfa = S.cfa_result is not None and S.cfa_result["success"]
            if has_cfa:
                st.info(f"Full EFA + CFA Word report (.docx)  **Cost: {REPORT_COST} credit**")
                if current_credits < REPORT_COST:
                    st.warning("⚠️ Insufficient credits.")
                    st.link_button("Top up →", "https://x.com/bayantx360", use_container_width=True)
                else:
                    if st.button(f"🔨 Build Word Report ({REPORT_COST} cr)", use_container_width=True, key="build_rpt"):
                        with st.spinner("Compiling Word report…"):
                            S.report_docx = generate_docx_report(
                                original_df=S.df_original, cleaned_df=S.df_working,
                                suitability=S.suitability, efa_result=S.efa_result,
                                diagnostics=S.diagnostics, dropped_vars=S.dropped_vars,
                                cfa_result=S.cfa_result, fit_assessment=S.fit_assessment,
                                cfa_thresholds=cfa_thresholds, synthetic_validation=S.syn_validation,
                                model_str=S.cfa_result.get("model_str",""))
                        handle_credit_deduction(REPORT_COST, app="EFActor", action="DOCX Report Export")
                        st.success(f"✓ Report ready. {REPORT_COST} credit deducted.")
                    if S.report_docx:
                        st.download_button("⬇ efactor_report.docx", data=S.report_docx,
                                           file_name="efactor_report.docx",
                                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                           use_container_width=True)
            else:
                st.warning("Run CFA (Step 5) to enable report generation.")

        st.markdown("---")
        st.markdown("##### 📦 Full Export Bundle (.zip)")
        zip_rows = max(len(S.df_working), len(syn_export) if syn_export is not None else 0)
        zip_cost = export_credit_cost(zip_rows) + (REPORT_COST if S.report_docx else 0)
        st.info(f"Cleaned data + synthetic data + Word report + loadings CSV + model spec  **Cost: {zip_cost} credits**")
        if current_credits < zip_cost:
            st.warning(f"⚠️ Need {zip_cost} credits ({current_credits} remaining).")
            st.link_button("Top up credits →", "https://x.com/bayantx360", use_container_width=True)
        else:
            if st.button(f"⬇ Build & Download ZIP ({zip_cost} cr)", use_container_width=True, key="dl_zip_btn"):
                zip_buffer = io.BytesIO()
                with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
                    zf.writestr("cleaned_data.csv", S.df_working.to_csv(index=False))
                    if syn_export is not None: zf.writestr(f"synthetic_{syn_label}.csv", syn_export.to_csv(index=False))
                    if S.report_docx: zf.writestr("efa_cfa_report.docx", S.report_docx)
                    if S.cfa_result and S.cfa_result.get("model_str"): zf.writestr("cfa_model.txt", S.cfa_result["model_str"])
                    if S.efa_result:
                        zf.writestr("efa_loadings.csv", S.efa_result["loadings"].round(4).to_csv())
                        zf.writestr("efa_communalities.csv", S.efa_result["communalities"].round(4).to_frame().to_csv())
                zip_buffer.seek(0)
                handle_credit_deduction(zip_cost, app="EFActor", action="Full ZIP Export")
                ts_str = datetime.now().strftime("%Y%m%d_%H%M")
                st.download_button("⬇ efactor_bundle.zip", data=zip_buffer.getvalue(),
                                   file_name=f"efactor_bundle_{ts_str}.zip", mime="application/zip", key="dl_zip_actual")
                st.success(f"✓ {zip_cost} credits deducted.")

st.markdown("---")
st.markdown('<div style="text-align:center;font-family:\'DM Mono\',monospace;font-size:0.68rem;color:var(--muted);padding:10px 0;">🔬 EFActor · Psychometric Analysis Platform · Bayantx360 Suite</div>', unsafe_allow_html=True)
